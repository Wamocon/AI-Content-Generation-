"""
Google Drive and Sheets Services for FIAE AI Content Factory
"""

import os
import json
import io
from typing import Dict, Any, List, Optional
from datetime import datetime
from loguru import logger

try:
    from google.oauth2.service_account import Credentials
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload, MediaIoBaseUpload
    from googleapiclient.errors import HttpError
    GOOGLE_APIS_AVAILABLE = True
except ImportError:
    GOOGLE_APIS_AVAILABLE = False
    logger.warning("Google APIs not available - install google-api-python-client")


class GoogleDriveService:
    """Google Drive service for file operations"""
    
    def __init__(self):
        self.service = None
        self.credentials = None
        self._initialize_service()
    
    def _initialize_service(self):
        """Initialize Google Drive service with personal account preference"""
        if not GOOGLE_APIS_AVAILABLE:
            logger.error("Google APIs not available")
            return
        
        try:
            # PRIORITY 1: Try personal Google account first
            personal_account_enabled = os.getenv('PERSONAL_GOOGLE_ACCOUNT_ENABLED', 'True').lower() == 'true'
            
            if personal_account_enabled:
                try:
                    # Import personal Google Drive service
                    import sys
                    sys.path.append(os.getcwd())
                    
                    from personal_google_drive_service import PersonalGoogleDriveService
                    
                    personal_service = PersonalGoogleDriveService(
                        credentials_file="personal_credentials.json",
                        token_file="personal_google_token.pickle"
                    )
                    
                    if personal_service.is_authenticated():
                        # Use personal account credentials
                        self.credentials = personal_service.credentials
                        self.service = build('drive', 'v3', credentials=self.credentials)
                        logger.info("[OK] Google Drive service initialized with PERSONAL account")
                        return
                    else:
                        logger.error("Personal Google account not authenticated. Please run authentication first.")
                        return
                        
                except Exception as e:
                    logger.error(f"Personal Google account failed: {e}. Please check your personal_credentials.json file.")
                    return
            
            # PRIORITY 2: Fallback to service account (DISABLED - GCP NOT REQUIRED)
            logger.error("Service account fallback disabled. Please use personal Google account authentication.")
            return
                
        except Exception as e:
            logger.error(f"Failed to initialize Google Drive service: {e}")
            self.service = None
    
    def list_files_in_folder(self, folder_id: str) -> List[Dict[str, Any]]:
        """List files in a Google Drive folder"""
        if not self.service:
            logger.error("Google Drive service not available")
            return []
        
        try:
            results = self.service.files().list(
                q=f"'{folder_id}' in parents and trashed=false",
                fields="files(id,name,mimeType,size,createdTime,modifiedTime)"
            ).execute()
            
            files = results.get('files', [])
            logger.info(f"Found {len(files)} files in folder {folder_id}")
            return files
            
        except Exception as e:
            logger.error(f"Error listing files in folder {folder_id}: {e}")
            return []
    
    def get_file_content(self, file_id: str) -> str:
        """Get content of a Google Drive file"""
        if not self.service:
            logger.error("Google Drive service not available")
            return ""
        
        try:
            # Get file metadata first
            file_metadata = self.service.files().get(fileId=file_id).execute()
            mime_type = file_metadata.get('mimeType', '')
            file_name = file_metadata.get('name', '')
            logger.info(f"[DEBUG] File: {file_name}, MIME Type: {mime_type}")
            
            # Handle different file types
            if 'wordprocessingml' in mime_type or mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                # Word documents (.docx) - download and extract text
                request = self.service.files().get_media(fileId=file_id)
                file_content = io.BytesIO()
                downloader = MediaIoBaseDownload(file_content, request)
                done = False
                while done is False:
                    status, done = downloader.next_chunk()
                
                # Extract text from .docx file
                try:
                    from docx import Document
                    file_content.seek(0)
                    doc = Document(file_content)
                    
                    # Extract all text from paragraphs
                    full_text = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text.strip())
                    
                    # Also extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    full_text.append(cell.text.strip())
                    
                    extracted_content = '\n\n'.join(full_text)
                    logger.info(f"[SUCCESS] Extracted {len(extracted_content)} characters from .docx file")
                    return extracted_content
                    
                except ImportError:
                    logger.error("python-docx not installed. Install with: pip install python-docx")
                    return f"Word document content for file {file_id} (python-docx not available)"
                except Exception as e:
                    logger.error(f"Error extracting text from .docx file: {e}")
                    return f"Word document content for file {file_id} (extraction failed: {str(e)})"
            elif 'document' in mime_type and 'google-apps' in mime_type:
                # Google Docs - export as plain text
                content = self.service.files().export(
                    fileId=file_id,
                    mimeType='text/plain'
                ).execute()
                return content.decode('utf-8')
            else:
                logger.warning(f"Unsupported file type: {mime_type}")
                return ""
                
        except Exception as e:
            logger.error(f"Error getting file content for {file_id}: {e}")
            return ""
    
    def save_comprehensive_content_to_review_folder(
        self, 
        original_filename: str, 
        content: Dict[str, Any], 
        source_file_id: str
    ) -> Dict[str, Any]:
        """Save comprehensive content to review folder"""
        if not self.service:
            return {"success": False, "error": "Google Drive service not available"}
        
        try:
            # Create folder for this document
            folder_name = f"FIAE_Generated_{original_filename}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            # Import settings to get the correct folder ID
            from app.config import settings
            
            folder_metadata = {
                'name': folder_name,
                'mimeType': 'application/vnd.google-apps.folder',
                'parents': [settings.google_drive_review_folder_id]
            }
            
            folder = self.service.files().create(
                body=folder_metadata,
                fields='id'
            ).execute()
            
            folder_id = folder.get('id')
            created_files = []
            
            # Save each content type as a separate file
            content_types = ['knowledge_analysis', 'use_case_text', 'quiz_text', 'powerpoint_structure', 'google_slides_content', 'trainer_script']
            
            for content_type in content_types:
                if content_type in content:
                    file_metadata = {
                        'name': f"{content_type}_{original_filename}.txt",
                        'parents': [folder_id]
                    }
                    
                    # Create a temporary file-like object for the content
                    content_data = content[content_type]
                    
                    # Handle both string and list content
                    if isinstance(content_data, list):
                        # If it's a list, join it into a string
                        content_text = '\n'.join(str(item) for item in content_data)
                    elif isinstance(content_data, dict):
                        # If it's a dict, convert to readable format
                        content_text = str(content_data)
                    else:
                        # If it's already a string, use it directly
                        content_text = str(content_data)
                    
                    content_bytes = content_text.encode('utf-8')
                    media = MediaIoBaseUpload(
                        io.BytesIO(content_bytes),
                        mimetype='text/plain'
                    )
                    
                    file = self.service.files().create(
                        body=file_metadata,
                        media_body=media,
                        fields='id'
                    ).execute()
                    
                    created_files.append({
                        'name': file_metadata['name'],
                        'id': file.get('id'),
                        'type': content_type
                    })
            
            return {
                "success": True,
                "enabler_folder_id": folder_id,
                "folder_name": folder_name,
                "created_files": created_files,
                "total_files_created": len(created_files)
            }
            
        except Exception as e:
            logger.error(f"Error saving content to review folder: {e}")
            return {"success": False, "error": str(e)}
    
    def detect_google_sheets_in_folder(self, folder_id: str) -> List[Dict[str, Any]]:
        """Detect Google Sheets in a folder"""
        if not self.service:
            return []
        
        try:
            results = self.service.files().list(
                q=f"'{folder_id}' in parents and mimeType='application/vnd.google-apps.spreadsheet' and trashed=false",
                fields="files(id,name,createdTime,modifiedTime)"
            ).execute()
            
            return results.get('files', [])
            
        except Exception as e:
            logger.error(f"Error detecting Google Sheets in folder {folder_id}: {e}")
            return []
    
    def _get_drive_storage_info(self) -> Dict[str, Any]:
        """Get Google Drive storage information"""
        if not self.service:
            return {"total_space": 0, "used_space": 0, "available_space": 0}
        
        try:
            about = self.service.about().get(fields="storageQuota").execute()
            quota = about.get('storageQuota', {})
            
            return {
                "total_space": int(quota.get('limit', 0)),
                "used_space": int(quota.get('usage', 0)),
                "available_space": int(quota.get('limit', 0)) - int(quota.get('usage', 0))
            }
            
        except Exception as e:
            logger.error(f"Error getting storage info: {e}")
            return {"total_space": 0, "used_space": 0, "available_space": 0}


class GoogleSheetsService:
    """Google Sheets service for tracking and analytics"""
    
    def __init__(self):
        self.service = None
        self.credentials = None
        self.spreadsheet_id = None
        self._initialize_service()
    
    def _initialize_service(self):
        """Initialize Google Sheets service"""
        if not GOOGLE_APIS_AVAILABLE:
            logger.error("Google APIs not available")
            return
        
        try:
            # Use personal Google account for Sheets access
            import sys
            sys.path.append(os.getcwd())
            
            from personal_google_drive_service import PersonalGoogleDriveService
            
            personal_service = PersonalGoogleDriveService(
                credentials_file="personal_credentials.json",
                token_file="personal_google_token.pickle"
            )
            
            if personal_service.is_authenticated():
                self.credentials = personal_service.credentials
                self.service = build('sheets', 'v4', credentials=self.credentials)
                self.spreadsheet_id = os.getenv('GOOGLE_SHEETS_ID', '1d87xmQNbWlNwtvRfhaWLSk2FkfTRVadKm94-ppaASbw')
                logger.info("[OK] Google Sheets service initialized with PERSONAL account")
            else:
                logger.error("Personal Google account not authenticated for Sheets access")
                self.service = None
                
        except Exception as e:
            logger.error(f"Failed to initialize Google Sheets service: {e}")
            self.service = None
    
    def add_processing_record(
        self, 
        job_id: str, 
        document_name: str, 
        status: str, 
        quality_score: float, 
        created_files: List[Dict[str, Any]]
    ) -> bool:
        """Add a processing record to Google Sheets"""
        if not self.service or not self.spreadsheet_id:
            logger.warning("Google Sheets service not available")
            return False
        
        try:
            # Prepare row data
            row_data = [
                job_id,
                document_name,
                status,
                quality_score,
                len(created_files),
                datetime.now().isoformat(),
                json.dumps(created_files)
            ]
            
            # Append to sheet
            body = {
                'values': [row_data]
            }
            
            result = self.service.spreadsheets().values().append(
                spreadsheetId=self.spreadsheet_id,
                range='Tabellenblatt1!A:M',  # Use correct German sheet name
                valueInputOption='RAW',
                body=body
            ).execute()
            
            logger.info(f"Added processing record for job {job_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error adding processing record: {e}")
            return False
    
    def get_processing_status(self) -> Dict[str, Any]:
        """Get processing status from Google Sheets"""
        if not self.service or not self.spreadsheet_id:
            return {
                "last_processing": None,
                "total_processed": 0,
                "pending_documents": 0,
                "failed_documents": 0,
                "processing_rate": "0%"
            }
        
        try:
            # Get all data from sheet
            result = self.service.spreadsheets().values().get(
                spreadsheetId=self.spreadsheet_id,
                range='Tabellenblatt1!A:M'
            ).execute()
            
            values = result.get('values', [])
            if len(values) <= 1:  # Only header or empty
                return {
                    "last_processing": None,
                    "total_processed": 0,
                    "pending_documents": 0,
                    "failed_documents": 0,
                    "processing_rate": "0%"
                }
            
            # Process data (skip header)
            data_rows = values[1:]
            total_processed = len([row for row in data_rows if len(row) > 2 and row[2] == 'completed'])
            failed_documents = len([row for row in data_rows if len(row) > 2 and row[2] == 'failed'])
            
            # Get last processing time
            last_processing = None
            if data_rows:
                last_row = data_rows[-1]
                if len(last_row) > 5:
                    last_processing = last_row[5]
            
            processing_rate = f"{(total_processed / len(data_rows) * 100):.1f}%" if data_rows else "0%"
            
            return {
                "last_processing": last_processing,
                "total_processed": total_processed,
                "pending_documents": len(data_rows) - total_processed - failed_documents,
                "failed_documents": failed_documents,
                "processing_rate": processing_rate
            }
            
        except Exception as e:
            logger.error(f"Error getting processing status: {e}")
            return {
                "last_processing": None,
                "total_processed": 0,
                "pending_documents": 0,
                "failed_documents": 0,
                "processing_rate": "0%"
            }
    
    def add_review_record(
        self,
        document_name: str,
        source_file_id: str,
        processing_status: str,
        output_folder_id: str,
        output_files: List[str],
        quality_score: float,
        processing_time: float = 0.0,
        gamma_pptx_file_id: str = ""
    ) -> bool:
        """Add document processing record for review tracking
        
        Google Sheets columns:
        A: Document_Name, B: Source_File_ID, C: Processing_Date, D: Processing_Status,
        E: Review_Status, F: Review_Date, G: Reviewer_Name, H: Output_Folder_ID,
        I: Output_Files_Created, J: Quality_Score, K: Error_Log, L: Processing_Time_Seconds, 
        M: Notes, N: Gamma_PPTX_File_ID
        """
        if not self.service or not self.spreadsheet_id:
            logger.warning("Google Sheets service not available")
            return False
        
        try:
            row_data = [
                document_name,  # A
                source_file_id,  # B
                datetime.now().isoformat(),  # C - Processing_Date
                processing_status,  # D - Processing_Status (completed/failed)
                "pending_review",  # E - Review_Status
                "",  # F - Review_Date (empty until reviewed)
                "",  # G - Reviewer_Name (empty until reviewed)
                output_folder_id,  # H - Output_Folder_ID
                ", ".join(output_files),  # I - Output_Files_Created
                f"{quality_score:.2f}",  # J - Quality_Score
                "",  # K - Error_Log
                f"{processing_time:.2f}",  # L - Processing_Time_Seconds
                "",  # M - Notes
                gamma_pptx_file_id  # N - Gamma_PPTX_File_ID
            ]
            
            # Check if sheet has headers, if not add them
            try:
                result = self.service.spreadsheets().values().get(
                    spreadsheetId=self.spreadsheet_id,
                    range='Sheet1!A1:N1'
                ).execute()
                
                if not result.get('values'):
                    # Add headers
                    headers = [[
                        "Document_Name", "Source_File_ID", "Processing_Date", "Processing_Status",
                        "Review_Status", "Review_Date", "Reviewer_Name", "Output_Folder_ID",
                        "Output_Files_Created", "Quality_Score", "Error_Log", "Processing_Time_Seconds", "Notes", "Gamma_PPTX_File_ID"
                    ]]
                    self.service.spreadsheets().values().update(
                        spreadsheetId=self.spreadsheet_id,
                        range='Sheet1!A1:N1',
                        valueInputOption='RAW',
                        body={'values': headers}
                    ).execute()
                    logger.info("Added headers to Google Sheets")
            except:
                pass  # Headers might already exist
            
            # Append review record
            body = {'values': [row_data]}
            result = self.service.spreadsheets().values().append(
                spreadsheetId=self.spreadsheet_id,
                range='Tabellenblatt1!A:M',
                valueInputOption='RAW',
                body=body
            ).execute()
            
            logger.info(f"✅ Added review record for {document_name}")
            return True
            
        except Exception as e:
            logger.error(f"Error adding review record: {e}")
            return False
    
    def get_pending_reviews(self) -> List[Dict[str, Any]]:
        """Get all documents pending review"""
        if not self.service or not self.spreadsheet_id:
            return []
        
        try:
            result = self.service.spreadsheets().values().get(
                spreadsheetId=self.spreadsheet_id,
                range='Tabellenblatt1!A:M'
            ).execute()
            
            values = result.get('values', [])
            if len(values) <= 1:  # Only header or empty
                return []
            
            # Skip header, filter for pending_review
            pending = []
            for i, row in enumerate(values[1:], start=2):  # Start at row 2 (after header)
                if len(row) >= 5 and row[4] == "pending_review":  # Column E
                    pending.append({
                        "row_number": i,
                        "document_name": row[0] if len(row) > 0 else "",
                        "source_file_id": row[1] if len(row) > 1 else "",
                        "processing_date": row[2] if len(row) > 2 else "",
                        "processing_status": row[3] if len(row) > 3 else "",
                        "review_status": row[4] if len(row) > 4 else "",
                        "output_folder_id": row[7] if len(row) > 7 else "",
                        "output_files": row[8] if len(row) > 8 else "",
                        "quality_score": float(row[9]) if len(row) > 9 and row[9] else 0.0,
                        "processing_time": float(row[11]) if len(row) > 11 and row[11] else 0.0
                    })
            
            logger.info(f"Found {len(pending)} documents pending review")
            return pending
            
        except Exception as e:
            logger.error(f"Error getting pending reviews: {e}")
            return []
    
    def update_review_status(
        self,
        document_name: str,
        review_status: str,  # "approved" or "rejected"
        reviewer_name: str = "User",
        notes: str = ""
    ) -> bool:
        """Update review status for a document"""
        if not self.service or not self.spreadsheet_id:
            logger.warning("Google Sheets service not available")
            return False
        
        try:
            # Find the row with this document name
            result = self.service.spreadsheets().values().get(
                spreadsheetId=self.spreadsheet_id,
                range='Tabellenblatt1!A:M'
            ).execute()
            
            values = result.get('values', [])
            if len(values) <= 1:
                return False
            
            # Find matching row
            for i, row in enumerate(values[1:], start=2):  # Start at row 2
                if len(row) > 0 and row[0] == document_name and len(row) >= 5 and row[4] == "pending_review":
                    # Update this row
                    update_range = f'Sheet1!E{i}:G{i}'  # Columns E, F, G
                    update_values = [[
                        review_status,  # E - Review_Status
                        datetime.now().isoformat(),  # F - Review_Date
                        reviewer_name  # G - Reviewer_Name
                    ]]
                    
                    # Also update notes if provided
                    if notes:
                        notes_range = f'Sheet1!M{i}'
                        self.service.spreadsheets().values().update(
                            spreadsheetId=self.spreadsheet_id,
                            range=notes_range,
                            valueInputOption='RAW',
                            body={'values': [[notes]]}
                        ).execute()
                    
                    self.service.spreadsheets().values().update(
                        spreadsheetId=self.spreadsheet_id,
                        range=update_range,
                        valueInputOption='RAW',
                        body={'values': update_values}
                    ).execute()
                    
                    logger.info(f"✅ Updated review status for {document_name} to {review_status}")
                    return True
            
            logger.warning(f"Document {document_name} not found in pending reviews")
            return False
            
        except Exception as e:
            logger.error(f"Error updating review status: {e}")
            return False
    
    def get_review_statistics(self) -> Dict[str, Any]:
        """Get review statistics"""
        if not self.service or not self.spreadsheet_id:
            return {
                "total_documents": 0,
                "pending_review": 0,
                "approved": 0,
                "rejected": 0
            }
        
        try:
            result = self.service.spreadsheets().values().get(
                spreadsheetId=self.spreadsheet_id,
                range='Tabellenblatt1!A:M'
            ).execute()
            
            values = result.get('values', [])
            if len(values) <= 1:
                return {"total_documents": 0, "pending_review": 0, "approved": 0, "rejected": 0}
            
            data_rows = values[1:]
            pending = len([r for r in data_rows if len(r) >= 5 and r[4] == "pending_review"])
            approved = len([r for r in data_rows if len(r) >= 5 and r[4] == "approved"])
            rejected = len([r for r in data_rows if len(r) >= 5 and r[4] == "rejected"])
            
            return {
                "total_documents": len(data_rows),
                "pending_review": pending,
                "approved": approved,
                "rejected": rejected
            }
            
        except Exception as e:
            logger.error(f"Error getting review statistics: {e}")
            return {"total_documents": 0, "pending_review": 0, "approved": 0, "rejected": 0}
    
    def get_pending_approvals(self) -> List[Dict[str, Any]]:
        """Get pending HITL approvals - alias for get_pending_reviews"""
        return self.get_pending_reviews()
    
    def get_approval_request(self, approval_id: str) -> Optional[Dict[str, Any]]:
        """Get specific approval request by row number"""
        pending = self.get_pending_reviews()
        for item in pending:
            if item.get("row_number") == int(approval_id):
                return item
        return None
    
    def update_approval_status(
        self, 
        approval_id: str, 
        status: str, 
        approved_by: str, 
        notes: str = ""
    ) -> bool:
        """Update approval status - alias for update_review_status"""
        # Get document name from approval_id (row number)
        pending = self.get_pending_reviews()
        for item in pending:
            if item.get("row_number") == int(approval_id):
                return self.update_review_status(
                    item['document_name'],
                    status,
                    approved_by,
                    notes
                )
        return False
    
    def get_approval_statistics(self) -> Dict[str, Any]:
        """Get approval statistics - alias for get_review_statistics"""
        return self.get_review_statistics()
    
    def update_job_status(
        self, 
        job_id: str, 
        status: str, 
        rejection_reason: str = "", 
        regenerate_phase: str = ""
    ) -> bool:
        """Update job status in sheets"""
        # Implementation for legacy compatibility
        return True
    
    def move_folder_to_done(self, folder_id: str, done_folder_id: str = "1yG_8-wBK1wfrEjzs5J_rKRRaHBpOFPoK") -> bool:
        """Move output folder from Review to Done folder
        
        Args:
            folder_id: ID of the folder to move
            done_folder_id: ID of the Done folder (default from config)
        
        Returns:
            True if successful, False otherwise
        """
        try:
            # Get DriveService from GoogleDriveService
            from app.config import settings
            
            # Import Google Drive service to access drive API
            credentials_path = os.getenv('GOOGLE_CREDENTIALS_PATH', 'credentials/wmc-automation-agents-e6ce75b3daa2.json')
            
            if not os.path.exists(credentials_path):
                logger.error(f"Credentials not found: {credentials_path}")
                return False
            
            from google.oauth2.service_account import Credentials
            from googleapiclient.discovery import build
            
            credentials = Credentials.from_service_account_file(
                credentials_path,
                scopes=['https://www.googleapis.com/auth/drive']
            )
            drive_service = build('drive', 'v3', credentials=credentials)
            
            # Get current parents
            file = drive_service.files().get(
                fileId=folder_id,
                fields='parents'
            ).execute()
            
            previous_parents = ",".join(file.get('parents', []))
            
            # Move folder: remove from current parents, add to done folder
            updated_file = drive_service.files().update(
                fileId=folder_id,
                addParents=done_folder_id,
                removeParents=previous_parents,
                fields='id, parents'
            ).execute()
            
            logger.info(f"✅ Moved folder {folder_id} to Done folder")
            return True
            
        except Exception as e:
            logger.error(f"Error moving folder to Done: {e}")
            return False
