"""
Advanced Document Processor with Semantic Chunking
Implements sophisticated document processing with multi-modal content analysis and adaptive processing strategies.
"""

import os
import json
import uuid
import asyncio
from datetime import datetime
from typing import Dict, Any, List, Optional, Tuple, Union
from dataclasses import dataclass
from enum import Enum
from loguru import logger
import numpy as np
import pandas as pd
from pathlib import Path
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    MAGIC_AVAILABLE = False
    magic = None
import PyPDF2
from docx import Document
from PIL import Image
import spacy
from sentence_transformers import SentenceTransformer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.metrics.pairwise import cosine_similarity
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

from app.config import settings


class DocumentType(Enum):
    """Document type enumeration."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    IMAGE = "image"
    UNKNOWN = "unknown"


class ChunkingStrategy(Enum):
    """Chunking strategy enumeration."""
    SEMANTIC = "semantic"
    STRUCTURAL = "structural"
    HYBRID = "hybrid"
    ADAPTIVE = "adaptive"


@dataclass
class DocumentMetadata:
    """Document metadata structure."""
    file_path: str
    file_name: str
    file_size: int
    file_type: DocumentType
    creation_date: datetime
    modification_date: datetime
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    language: Optional[str] = None
    encoding: Optional[str] = None


@dataclass
class SemanticChunk:
    """Semantic chunk structure."""
    chunk_id: str
    content: str
    chunk_type: str
    position: int
    start_char: int
    end_char: int
    word_count: int
    char_count: int
    semantic_embedding: Optional[np.ndarray] = None
    topics: List[str] = None
    entities: List[str] = None
    sentiment: Optional[float] = None
    importance_score: float = 0.0
    coherence_score: float = 0.0
    metadata: Dict[str, Any] = None


@dataclass
class ProcessingResult:
    """Document processing result structure."""
    success: bool
    document_metadata: DocumentMetadata
    chunks: List[SemanticChunk]
    processing_time: float
    quality_score: float
    error_message: Optional[str] = None
    processing_strategy: Optional[str] = None


class AdvancedDocumentProcessor:
    """
    Advanced Document Processor with semantic understanding and multi-modal analysis.
    
    Features:
    - Multi-format document support (PDF, DOCX, TXT, Images)
    - Semantic chunking with context preservation
    - Multi-modal content processing
    - Adaptive processing strategies
    - Document structure analysis
    - Language detection and processing
    - Entity extraction and topic modeling
    - Quality assessment and optimization
    """
    
    def __init__(self):
        """Initialize the advanced document processor."""
        self.temp_dir = settings.temp_dir
        self.max_file_size = settings.max_file_size_mb * 1024 * 1024  # Convert to bytes
        self.supported_formats = {'.pdf', '.docx', '.txt', '.png', '.jpg', '.jpeg', '.gif', '.bmp'}
        
        # Initialize NLP components
        self._initialize_nlp_models()
        self._initialize_embeddings()
        self._initialize_analyzers()
        
        # Create temp directory
        os.makedirs(self.temp_dir, exist_ok=True)
        
        logger.info("Advanced Document Processor initialized")
    
    def _initialize_nlp_models(self):
        """Initialize NLP models for text processing."""
        try:
            # Load spaCy model (try German first, fallback to English)
            try:
                self.nlp = spacy.load("de_core_news_sm")
                self.language = "de"
                logger.info("German spaCy model loaded")
            except OSError:
                try:
                    self.nlp = spacy.load("en_core_web_sm")
                    self.language = "en"
                    logger.info("English spaCy model loaded")
                except OSError:
                    # Fallback to basic model
                    self.nlp = spacy.blank("en")
                    self.language = "en"
                    logger.warning("Using basic spaCy model - install language models for better performance")
            
            # Initialize TF-IDF vectorizer
            self.tfidf_vectorizer = TfidfVectorizer(
                max_features=1000,
                stop_words='english',
                ngram_range=(1, 2)
            )
            
            # Initialize topic modeling
            self.topic_modeler = KMeans(n_clusters=5, random_state=42)
            
        except Exception as e:
            logger.error(f"Failed to initialize NLP models: {str(e)}")
            raise
    
    def _initialize_embeddings(self):
        """Initialize embedding model for semantic analysis."""
        try:
            self.embeddings_model = SentenceTransformer('all-MiniLM-L6-v2')
            logger.info("Sentence transformer model initialized")
        except Exception as e:
            logger.error(f"Failed to initialize embeddings: {str(e)}")
            raise
    
    def _initialize_analyzers(self):
        """Initialize content analyzers."""
        try:
            # Initialize file type detector
            self.file_detector = magic.Magic(mime=True)
            
            # Initialize image processors
            self.image_formats = {'.png', '.jpg', '.jpeg', '.gif', '.bmp'}
            
            logger.info("Content analyzers initialized")
            
        except Exception as e:
            logger.error(f"Failed to initialize analyzers: {str(e)}")
            raise
    
    async def process_document(
        self, 
        file_path: str,
        job_id: str,
        chunking_strategy: ChunkingStrategy = ChunkingStrategy.ADAPTIVE
    ) -> ProcessingResult:
        """
        Process a document with advanced semantic analysis.
        
        Args:
            file_path: Path to the document file
            job_id: Unique job identifier
            chunking_strategy: Strategy for chunking the document
        
        Returns:
            Processing result with semantic chunks and metadata
        """
        start_time = datetime.utcnow()
        
        try:
            logger.info(f"Processing document: {file_path} with strategy: {chunking_strategy.value}")
            
            # Validate file
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Extract document metadata
            metadata = await self._extract_document_metadata(file_path)
            
            # Determine optimal chunking strategy
            if chunking_strategy == ChunkingStrategy.ADAPTIVE:
                chunking_strategy = await self._determine_optimal_strategy(metadata)
            
            # Extract content based on document type
            content = await self._extract_document_content(file_path, metadata.file_type)
            
            # Perform semantic chunking
            chunks = await self._perform_semantic_chunking(
                content, 
                metadata, 
                chunking_strategy, 
                job_id
            )
            
            # Analyze chunks for quality and coherence
            analyzed_chunks = await self._analyze_chunks(chunks, job_id)
            
            # Calculate overall quality score
            quality_score = self._calculate_document_quality(analyzed_chunks, metadata)
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            result = ProcessingResult(
                success=True,
                document_metadata=metadata,
                chunks=analyzed_chunks,
                processing_time=processing_time,
                quality_score=quality_score,
                processing_strategy=chunking_strategy.value
            )
            
            logger.info(f"Document processing completed: {len(analyzed_chunks)} chunks, quality: {quality_score:.2f}")
            return result
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {str(e)}")
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            return ProcessingResult(
                success=False,
                document_metadata=DocumentMetadata(
                    file_path=file_path,
                    file_name=os.path.basename(file_path),
                    file_size=0,
                    file_type=DocumentType.UNKNOWN,
                    creation_date=datetime.utcnow(),
                    modification_date=datetime.utcnow()
                ),
                chunks=[],
                processing_time=processing_time,
                quality_score=0.0,
                error_message=str(e)
            )
    
    async def _extract_document_metadata(self, file_path: str) -> DocumentMetadata:
        """Extract comprehensive document metadata."""
        try:
            file_stat = os.stat(file_path)
            file_name = os.path.basename(file_path)
            file_extension = Path(file_path).suffix.lower()
            
            # Determine document type
            if file_extension == '.pdf':
                doc_type = DocumentType.PDF
            elif file_extension == '.docx':
                doc_type = DocumentType.DOCX
            elif file_extension == '.txt':
                doc_type = DocumentType.TXT
            elif file_extension in self.image_formats:
                doc_type = DocumentType.IMAGE
            else:
                doc_type = DocumentType.UNKNOWN
            
            # Get MIME type
            mime_type = self.file_detector.from_file(file_path)
            
            # Extract additional metadata based on file type
            page_count = None
            word_count = None
            language = None
            encoding = None
            
            if doc_type == DocumentType.PDF:
                page_count = await self._get_pdf_page_count(file_path)
            elif doc_type == DocumentType.DOCX:
                page_count, word_count = await self._get_docx_metadata(file_path)
            elif doc_type == DocumentType.TXT:
                encoding = await self._detect_text_encoding(file_path)
            
            return DocumentMetadata(
                file_path=file_path,
                file_name=file_name,
                file_size=file_stat.st_size,
                file_type=doc_type,
                creation_date=datetime.fromtimestamp(file_stat.st_ctime),
                modification_date=datetime.fromtimestamp(file_stat.st_mtime),
                page_count=page_count,
                word_count=word_count,
                language=language,
                encoding=encoding
            )
            
        except Exception as e:
            logger.error(f"Error extracting document metadata: {str(e)}")
            raise
    
    async def _get_pdf_page_count(self, file_path: str) -> int:
        """Get page count for PDF files."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        except Exception as e:
            logger.error(f"Error getting PDF page count: {str(e)}")
            return 0
    
    async def _get_docx_metadata(self, file_path: str) -> Tuple[int, int]:
        """Get metadata for DOCX files."""
        try:
            doc = Document(file_path)
            
            # Count pages (approximate)
            page_count = len(doc.paragraphs) // 20  # Rough estimate
            
            # Count words
            word_count = 0
            for paragraph in doc.paragraphs:
                word_count += len(paragraph.text.split())
            
            return page_count, word_count
            
        except Exception as e:
            logger.error(f"Error getting DOCX metadata: {str(e)}")
            return 0, 0
    
    async def _detect_text_encoding(self, file_path: str) -> str:
        """Detect text file encoding."""
        try:
            import chardet
            with open(file_path, 'rb') as file:
                raw_data = file.read(10000)  # Read first 10KB
                result = chardet.detect(raw_data)
                return result.get('encoding', 'utf-8')
        except Exception as e:
            logger.error(f"Error detecting text encoding: {str(e)}")
            return 'utf-8'
    
    async def _extract_document_content(
        self, 
        file_path: str, 
        doc_type: DocumentType
    ) -> str:
        """Extract content from different document types."""
        try:
            if doc_type == DocumentType.PDF:
                return await self._extract_pdf_content(file_path)
            elif doc_type == DocumentType.DOCX:
                return await self._extract_docx_content(file_path)
            elif doc_type == DocumentType.TXT:
                return await self._extract_text_content(file_path)
            elif doc_type == DocumentType.IMAGE:
                return await self._extract_image_content(file_path)
            else:
                raise ValueError(f"Unsupported document type: {doc_type}")
                
        except Exception as e:
            logger.error(f"Error extracting document content: {str(e)}")
            raise
    
    async def _extract_pdf_content(self, file_path: str) -> str:
        """Extract text content from PDF files."""
        try:
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text.strip():
                        content.append(page_text)
            
            return '\n\n'.join(content)
            
        except Exception as e:
            logger.error(f"Error extracting PDF content: {str(e)}")
            raise
    
    async def _extract_docx_content(self, file_path: str) -> str:
        """Extract text content from DOCX files."""
        try:
            doc = Document(file_path)
            content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))
            
            return '\n\n'.join(content)
            
        except Exception as e:
            logger.error(f"Error extracting DOCX content: {str(e)}")
            raise
    
    async def _extract_text_content(self, file_path: str) -> str:
        """Extract content from text files."""
        try:
            # Detect encoding
            encoding = await self._detect_text_encoding(file_path)
            
            with open(file_path, 'r', encoding=encoding) as file:
                return file.read()
                
        except Exception as e:
            logger.error(f"Error extracting text content: {str(e)}")
            raise
    
    async def _extract_image_content(self, file_path: str) -> str:
        """Extract content from image files (OCR would be implemented here)."""
        try:
            # For now, return basic image metadata
            # In production, this would use OCR (Tesseract, etc.)
            image = Image.open(file_path)
            return f"Image content: {image.size[0]}x{image.size[1]} pixels, mode: {image.mode}"
            
        except Exception as e:
            logger.error(f"Error extracting image content: {str(e)}")
            raise
    
    async def _determine_optimal_strategy(self, metadata: DocumentMetadata) -> ChunkingStrategy:
        """Determine optimal chunking strategy based on document characteristics."""
        try:
            # Analyze document characteristics
            if metadata.file_type == DocumentType.PDF and metadata.page_count and metadata.page_count > 10:
                return ChunkingStrategy.STRUCTURAL
            elif metadata.file_type == DocumentType.DOCX and metadata.word_count and metadata.word_count > 5000:
                return ChunkingStrategy.HYBRID
            elif metadata.file_type == DocumentType.TXT:
                return ChunkingStrategy.SEMANTIC
            else:
                return ChunkingStrategy.SEMANTIC
                
        except Exception as e:
            logger.error(f"Error determining optimal strategy: {str(e)}")
            return ChunkingStrategy.SEMANTIC
    
    async def _perform_semantic_chunking(
        self, 
        content: str, 
        metadata: DocumentMetadata,
        strategy: ChunkingStrategy,
        job_id: str
    ) -> List[SemanticChunk]:
        """Perform semantic chunking based on the selected strategy."""
        try:
            if strategy == ChunkingStrategy.SEMANTIC:
                return await self._semantic_chunking(content, metadata, job_id)
            elif strategy == ChunkingStrategy.STRUCTURAL:
                return await self._structural_chunking(content, metadata, job_id)
            elif strategy == ChunkingStrategy.HYBRID:
                return await self._hybrid_chunking(content, metadata, job_id)
            else:
                return await self._semantic_chunking(content, metadata, job_id)
                
        except Exception as e:
            logger.error(f"Error performing semantic chunking: {str(e)}")
            raise
    
    async def _semantic_chunking(
        self, 
        content: str, 
        metadata: DocumentMetadata,
        job_id: str
    ) -> List[SemanticChunk]:
        """Perform semantic-based chunking."""
        try:
            # Split content into sentences
            doc = self.nlp(content)
            sentences = [sent.text for sent in doc.sents if sent.text.strip()]
            
            # Create embeddings for sentences
            embeddings = self.embeddings_model.encode(sentences)
            
            # Cluster sentences based on semantic similarity
            if len(sentences) > 1:
                # Use K-means clustering
                n_clusters = min(max(len(sentences) // 5, 2), 10)
                clusterer = KMeans(n_clusters=n_clusters, random_state=42)
                cluster_labels = clusterer.fit_predict(embeddings)
            else:
                cluster_labels = [0]
            
            # Group sentences by cluster
            clusters = {}
            for i, (sentence, label) in enumerate(zip(sentences, cluster_labels)):
                if label not in clusters:
                    clusters[label] = []
                clusters[label].append((i, sentence))
            
            # Create chunks from clusters
            chunks = []
            chunk_id = 0
            
            for cluster_id, cluster_sentences in clusters.items():
                # Sort sentences by original order
                cluster_sentences.sort(key=lambda x: x[0])
                
                # Combine sentences into chunk
                chunk_content = ' '.join([sent[1] for sent in cluster_sentences])
                
                # Calculate position
                start_char = content.find(chunk_content)
                end_char = start_char + len(chunk_content)
                
                chunk = SemanticChunk(
                    chunk_id=f"{job_id}_semantic_{chunk_id}",
                    content=chunk_content,
                    chunk_type="semantic",
                    position=chunk_id,
                    start_char=start_char,
                    end_char=end_char,
                    word_count=len(chunk_content.split()),
                    char_count=len(chunk_content),
                    metadata={
                        "cluster_id": cluster_id,
                        "sentence_count": len(cluster_sentences),
                        "strategy": "semantic"
                    }
                )
                
                chunks.append(chunk)
                chunk_id += 1
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error in semantic chunking: {str(e)}")
            raise
    
    async def _structural_chunking(
        self, 
        content: str, 
        metadata: DocumentMetadata,
        job_id: str
    ) -> List[SemanticChunk]:
        """Perform structure-based chunking."""
        try:
            # Split by paragraphs first
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            
            chunks = []
            chunk_id = 0
            current_chunk = ""
            current_start = 0
            
            for paragraph in paragraphs:
                # Check if adding this paragraph would exceed optimal chunk size
                if current_chunk and len(current_chunk) + len(paragraph) > 1000:
                    # Create chunk from current content
                    chunk = SemanticChunk(
                        chunk_id=f"{job_id}_structural_{chunk_id}",
                        content=current_chunk.strip(),
                        chunk_type="structural",
                        position=chunk_id,
                        start_char=current_start,
                        end_char=current_start + len(current_chunk),
                        word_count=len(current_chunk.split()),
                        char_count=len(current_chunk),
                        metadata={
                            "paragraph_count": current_chunk.count('\n\n') + 1,
                            "strategy": "structural"
                        }
                    )
                    chunks.append(chunk)
                    
                    # Start new chunk
                    chunk_id += 1
                    current_chunk = paragraph
                    current_start += len(current_chunk)
                else:
                    current_chunk += '\n\n' + paragraph if current_chunk else paragraph
            
            # Add final chunk
            if current_chunk:
                chunk = SemanticChunk(
                    chunk_id=f"{job_id}_structural_{chunk_id}",
                    content=current_chunk.strip(),
                    chunk_type="structural",
                    position=chunk_id,
                    start_char=current_start,
                    end_char=current_start + len(current_chunk),
                    word_count=len(current_chunk.split()),
                    char_count=len(current_chunk),
                    metadata={
                        "paragraph_count": current_chunk.count('\n\n') + 1,
                        "strategy": "structural"
                    }
                )
                chunks.append(chunk)
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error in structural chunking: {str(e)}")
            raise
    
    async def _hybrid_chunking(
        self, 
        content: str, 
        metadata: DocumentMetadata,
        job_id: str
    ) -> List[SemanticChunk]:
        """Perform hybrid chunking combining semantic and structural approaches."""
        try:
            # First, do structural chunking
            structural_chunks = await self._structural_chunking(content, metadata, job_id)
            
            # Then, refine with semantic analysis
            refined_chunks = []
            chunk_id = 0
            
            for chunk in structural_chunks:
                # If chunk is too large, split it semantically
                if chunk.char_count > 2000:
                    # Split into sentences
                    doc = self.nlp(chunk.content)
                    sentences = [sent.text for sent in doc.sents if sent.text.strip()]
                    
                    # Create sub-chunks
                    sub_chunks = []
                    current_sub_chunk = ""
                    
                    for sentence in sentences:
                        if len(current_sub_chunk) + len(sentence) > 1000 and current_sub_chunk:
                            sub_chunks.append(current_sub_chunk.strip())
                            current_sub_chunk = sentence
                        else:
                            current_sub_chunk += ' ' + sentence if current_sub_chunk else sentence
                    
                    if current_sub_chunk:
                        sub_chunks.append(current_sub_chunk.strip())
                    
                    # Create refined chunks
                    for i, sub_content in enumerate(sub_chunks):
                        refined_chunk = SemanticChunk(
                            chunk_id=f"{job_id}_hybrid_{chunk_id}",
                            content=sub_content,
                            chunk_type="hybrid",
                            position=chunk_id,
                            start_char=chunk.start_char + chunk.content.find(sub_content),
                            end_char=chunk.start_char + chunk.content.find(sub_content) + len(sub_content),
                            word_count=len(sub_content.split()),
                            char_count=len(sub_content),
                            metadata={
                                "parent_chunk": chunk.chunk_id,
                                "sub_chunk_index": i,
                                "strategy": "hybrid"
                            }
                        )
                        refined_chunks.append(refined_chunk)
                        chunk_id += 1
                else:
                    # Keep original chunk but update metadata
                    chunk.chunk_id = f"{job_id}_hybrid_{chunk_id}"
                    chunk.chunk_type = "hybrid"
                    chunk.position = chunk_id
                    chunk.metadata["strategy"] = "hybrid"
                    refined_chunks.append(chunk)
                    chunk_id += 1
            
            return refined_chunks
            
        except Exception as e:
            logger.error(f"Error in hybrid chunking: {str(e)}")
            raise
    
    async def _analyze_chunks(self, chunks: List[SemanticChunk], job_id: str) -> List[SemanticChunk]:
        """Analyze chunks for semantic properties and quality."""
        try:
            analyzed_chunks = []
            
            for chunk in chunks:
                # Extract topics using TF-IDF
                topics = await self._extract_topics(chunk.content)
                chunk.topics = topics
                
                # Extract entities using spaCy
                entities = await self._extract_entities(chunk.content)
                chunk.entities = entities
                
                # Calculate sentiment
                sentiment = await self._calculate_sentiment(chunk.content)
                chunk.sentiment = sentiment
                
                # Calculate importance score
                importance = await self._calculate_importance_score(chunk)
                chunk.importance_score = importance
                
                # Calculate coherence score
                coherence = await self._calculate_coherence_score(chunk)
                chunk.coherence_score = coherence
                
                # Create semantic embedding
                embedding = self.embeddings_model.encode([chunk.content])[0]
                chunk.semantic_embedding = embedding
                
                analyzed_chunks.append(chunk)
            
            return analyzed_chunks
            
        except Exception as e:
            logger.error(f"Error analyzing chunks: {str(e)}")
            return chunks
    
    async def _extract_topics(self, content: str) -> List[str]:
        """Extract topics from content using TF-IDF."""
        try:
            # Simple topic extraction using TF-IDF
            words = content.lower().split()
            
            # Filter out common words
            stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by'}
            filtered_words = [word for word in words if len(word) > 3 and word not in stop_words]
            
            # Get most frequent words as topics
            word_freq = {}
            for word in filtered_words:
                word_freq[word] = word_freq.get(word, 0) + 1
            
            # Return top 5 topics
            topics = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:5]
            return [topic[0] for topic in topics]
            
        except Exception as e:
            logger.error(f"Error extracting topics: {str(e)}")
            return []
    
    async def _extract_entities(self, content: str) -> List[str]:
        """Extract entities using spaCy."""
        try:
            doc = self.nlp(content)
            entities = []
            
            for ent in doc.ents:
                if ent.label_ in ['PERSON', 'ORG', 'GPE', 'PRODUCT', 'EVENT']:
                    entities.append(ent.text)
            
            return list(set(entities))  # Remove duplicates
            
        except Exception as e:
            logger.error(f"Error extracting entities: {str(e)}")
            return []
    
    async def _calculate_sentiment(self, content: str) -> float:
        """Calculate sentiment score for content."""
        try:
            # Simple sentiment analysis based on positive/negative words
            positive_words = {'good', 'great', 'excellent', 'amazing', 'wonderful', 'fantastic', 'positive', 'success'}
            negative_words = {'bad', 'terrible', 'awful', 'horrible', 'negative', 'failure', 'problem', 'issue'}
            
            words = content.lower().split()
            positive_count = sum(1 for word in words if word in positive_words)
            negative_count = sum(1 for word in words if word in negative_words)
            
            if positive_count + negative_count == 0:
                return 0.0
            
            sentiment = (positive_count - negative_count) / (positive_count + negative_count)
            return max(-1.0, min(1.0, sentiment))  # Clamp between -1 and 1
            
        except Exception as e:
            logger.error(f"Error calculating sentiment: {str(e)}")
            return 0.0
    
    async def _calculate_importance_score(self, chunk: SemanticChunk) -> float:
        """Calculate importance score for a chunk."""
        try:
            score = 0.0
            
            # Length factor
            if chunk.word_count > 50:
                score += 0.2
            elif chunk.word_count > 20:
                score += 0.1
            
            # Entity factor
            if chunk.entities and len(chunk.entities) > 0:
                score += 0.2
            
            # Topic factor
            if chunk.topics and len(chunk.topics) > 0:
                score += 0.2
            
            # Position factor (earlier chunks might be more important)
            if chunk.position < 3:
                score += 0.2
            
            # Coherence factor
            if chunk.coherence_score > 0.7:
                score += 0.2
            
            return min(score, 1.0)
            
        except Exception as e:
            logger.error(f"Error calculating importance score: {str(e)}")
            return 0.5
    
    async def _calculate_coherence_score(self, chunk: SemanticChunk) -> float:
        """Calculate coherence score for a chunk."""
        try:
            # Simple coherence based on sentence structure and flow
            sentences = chunk.content.split('.')
            if len(sentences) < 2:
                return 0.5
            
            # Check for transition words
            transition_words = {'however', 'therefore', 'furthermore', 'moreover', 'additionally', 'consequently'}
            transition_count = sum(1 for word in chunk.content.lower().split() if word in transition_words)
            
            # Check for pronoun references
            pronouns = {'it', 'this', 'that', 'these', 'those', 'he', 'she', 'they'}
            pronoun_count = sum(1 for word in chunk.content.lower().split() if word in pronouns)
            
            # Calculate coherence
            coherence = min(1.0, (transition_count + pronoun_count) / len(sentences))
            return coherence
            
        except Exception as e:
            logger.error(f"Error calculating coherence score: {str(e)}")
            return 0.5
    
    def _calculate_document_quality(
        self, 
        chunks: List[SemanticChunk], 
        metadata: DocumentMetadata
    ) -> float:
        """Calculate overall document quality score."""
        try:
            if not chunks:
                return 0.0
            
            # Calculate average scores
            avg_importance = np.mean([chunk.importance_score for chunk in chunks])
            avg_coherence = np.mean([chunk.coherence_score for chunk in chunks])
            
            # Calculate coverage (how much of the document is chunked)
            total_chars = sum(chunk.char_count for chunk in chunks)
            coverage = min(1.0, total_chars / (metadata.file_size / 10))  # Rough estimate
            
            # Calculate diversity (number of unique topics)
            all_topics = []
            for chunk in chunks:
                if chunk.topics:
                    all_topics.extend(chunk.topics)
            topic_diversity = len(set(all_topics)) / max(len(all_topics), 1)
            
            # Combine factors
            quality = (
                avg_importance * 0.3 +
                avg_coherence * 0.3 +
                coverage * 0.2 +
                topic_diversity * 0.2
            )
            
            return min(quality, 1.0)
            
        except Exception as e:
            logger.error(f"Error calculating document quality: {str(e)}")
            return 0.5
    
    async def get_processing_statistics(self) -> Dict[str, Any]:
        """Get processing statistics and performance metrics."""
        try:
            return {
                "processor_status": "active",
                "supported_formats": list(self.supported_formats),
                "max_file_size_mb": settings.max_file_size_mb,
                "nlp_language": self.language,
                "embeddings_model": "all-MiniLM-L6-v2",
                "chunking_strategies": [strategy.value for strategy in ChunkingStrategy],
                "temp_directory": self.temp_dir
            }
            
        except Exception as e:
            logger.error(f"Error getting processing statistics: {str(e)}")
            return {"error": str(e)}

