#!/usr/bin/env python3
"""
PHASE 1: DiTeLe STANDARD SCENARIO GENERATION
=============================================
🎯 Focus: Generate educational scenarios following DiTeLe standard structure
📚 Coverage: Multiple problem-solution pairs covering all major topics
✅ Quality: Detailed, beginner-friendly, trainer-ready
📄 Output: Professional Word documents with standardized structure

DiTeLe STRUCTURE:
1. Themenliste (Topic List)
2. Lernziele (Learning Objectives - upfront)
3. Theoretische Grundlagen (Theoretical Foundation - 800+ words)  
4. Ausgangslage (Starting Situation - realistic context)
5. Problem-Lösungs-Paare (Multiple Problem-Solution Pairs - one per topic)
6. Lernziel-Checkliste (Learning Objectives Checklist - verification)

TEST MODE:
- Set TEST_MODE = True to process only 2 documents for quality review
- Set TEST_MODE = False for full batch processing
- Always test first before running full automation!

Workflow:
1. Analyze source document (identify topics, complexity, scope)
2. Generate DiTeLe-compliant scenario with multiple problem-solution pairs
3. Validate structure completeness
4. Create professional Word document with proper formatting
5. Upload to Google Drive for review
"""

# ========== CONFIGURATION ==========
TEST_MODE = True  # Set to False after testing first 2 documents
MAX_TEST_DOCUMENTS = 2  # Number of documents to process in test mode

import asyncio
import os
import sys
from datetime import datetime
from typing import Dict, Any, List, Optional
from loguru import logger
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googleapiclient.http import MediaIoBaseUpload
import io
import re

load_dotenv()
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app.config import settings
from app.services.google_services import GoogleDriveService
from app.services.intelligent_gemini_service import IntelligentGeminiService
from app.services.document_analyzer import DocumentAnalyzer, AnalysisResult, Topic

# Configure logging
logger.remove()
logger.add(
    sys.stdout,
    format="<green>{time:YYYY-MM-DD HH:mm:ss}</green> | <level>{level: <8}</level> | <level>{message}</level>",
    level="INFO"
)

# Initialize services
google_drive_service = GoogleDriveService()
intelligent_gemini = IntelligentGeminiService()
document_analyzer = DocumentAnalyzer(gemini_service=intelligent_gemini.gemini_service)


async def analyze_document_for_ditele(document_content: str, doc_name: str) -> AnalysisResult:
    """
    🧠 Analyze document for DiTeLe scenario generation
    
    Identifies major topics that will become individual problem-solution pairs
    
    Returns:
        AnalysisResult with topics, complexity, and content requirements
    """
    logger.info(f"🧠 Analysiere Dokument: {doc_name}")
    
    # Use AI-powered analysis
    analysis_result = await document_analyzer.analyze_document(
        document_content=document_content,
        document_name=doc_name,
        use_ai=True
    )
    
    # Log analysis results
    logger.info(f"📊 Analyse-Ergebnisse:")
    logger.info(f"   📝 Wörter: {len(document_content.split()):,}")
    logger.info(f"   📑 Themen identifiziert: {len(analysis_result.topics)}")
    logger.info(f"   🎯 Komplexität: {analysis_result.complexity_score:.1f}/10")
    logger.info(f"   💡 Problem-Lösungs-Paare: {min(len(analysis_result.topics), 7)}")
    
    return analysis_result


def _generate_problem_solution_template(topics: List[str], start_num: int = 1) -> str:
    """Helper to generate problem-solution pair templates for each topic"""
    template = ""
    for i, topic in enumerate(topics, start_num):
        template += f"""
PROBLEM {i}: {topic}
───────────────────────────────────────────────────────────────

Situation:
[Konkreter Kontext im Unternehmen]

Aufgabe:
[Was soll genau getan werden?]

Randbedingungen:
[Vorgaben, Constraints, Anforderungen]

Erwartete Ergebnisse:
[Was wird am Ende erwartet?]

═══════════════════════════════════════════════════════════════

LÖSUNG {i}: {topic}
───────────────────────────────────────────────────────────────

Schritt 1: [Erster Schritt]
[Konkrete Anleitung]
[Beispiel-Code wenn noetig]

Schritt 2: [Zweiter Schritt]
[Konkrete Anleitung]
[Tools/Ressourcen die benoetigt werden]

Schritt 3: [Dritter Schritt]
[Konkrete Anleitung]
[Tools/Ressourcen die benötigt werden]

Schritt 4-N: [Weitere Schritte bis zur vollständigen Lösung]
[ALLE Schritte müssen erklärt sein]
[Anfänger müssen JEDEN Schritt nachvollziehen können]

Ergebnis:
[Was ist das finale Ergebnis nach allen Schritten?]
[Wie sieht die Lösung konkret aus?]

Erklärung:
[WARUM funktioniert diese Lösung?]
[Welche Prinzipien/Konzepte stecken dahinter?]

Alternative Ansätze:
[Gibt es andere Lösungswege?]
[Vor- und Nachteile diskutieren]

Häufige Fehler:
[Top 2-3 Fehler die Anfänger machen]
[Wie vermeidet/behebt man diese?]

═══════════════════════════════════════════════════════════════
"""
    return template.strip()


def _cleanup_batch_content(content: str, expected_start: int, expected_end: int) -> str:
    """
    Clean up batch content to ensure numbering consistency.
    
    Fixes:
    - Renumbers problems to expected sequential numbers
    - Removes forbidden terms (bot, KI, AI)
    - Ensures consistency
    """
    import re
    
    # Remove forbidden terms
    forbidden_patterns = [
        (r'\b[Bb]ot\b', '[System]'),
        (r'\b[Kk][Ii]\b', 'System'),
        (r'\b[Aa][Ii]\b', 'System'),
        (r'[Qq]uality [Ss]core', 'Qualität'),
        (r'[Qq]ualitätsscore', 'Qualität'),
    ]
    
    cleaned = content
    for pattern, replacement in forbidden_patterns:
        cleaned = re.sub(pattern, replacement, cleaned)
    
    # Find all PROBLEM headers with their numbers
    problem_pattern = r'PROBLEM\s+(\d+):'
    problems_found = list(re.finditer(problem_pattern, cleaned))
    
    if problems_found:
        # Renumber problems sequentially
        problem_mapping = {}
        expected_num = expected_start
        
        for match in problems_found:
            old_num = match.group(1)
            if old_num not in problem_mapping:
                problem_mapping[old_num] = expected_num
                expected_num += 1
        
        # Apply renumbering (replace from end to start to avoid conflicts)
        for match in reversed(problems_found):
            old_num = match.group(1)
            new_num = problem_mapping[old_num]
            old_text = f"PROBLEM {old_num}:"
            new_text = f"PROBLEM {new_num}:"
            
            # Replace only at the specific position
            start_pos = match.start()
            end_pos = match.end()
            cleaned = cleaned[:start_pos] + new_text + cleaned[end_pos:]
        
        logger.info(f"      🔧 Renumbered {len(problem_mapping)} problems: {list(problem_mapping.values())}")
    
    # Also renumber LÖSUNG headers
    solution_pattern = r'LÖSUNG\s+(\d+):'
    solutions_found = list(re.finditer(solution_pattern, cleaned))
    
    if solutions_found:
        problem_mapping = {}
        expected_num = expected_start
        
        for match in solutions_found:
            old_num = match.group(1)
            if old_num not in problem_mapping:
                problem_mapping[old_num] = expected_num
                expected_num += 1
        
        for match in reversed(solutions_found):
            old_num = match.group(1)
            new_num = problem_mapping[old_num]
            old_text = f"LÖSUNG {old_num}:"
            new_text = f"LÖSUNG {new_num}:"
            
            start_pos = match.start()
            end_pos = match.end()
            cleaned = cleaned[:start_pos] + new_text + cleaned[end_pos:]
    
    return cleaned


def _extract_context_safely(content: str) -> tuple:
    """
    Extract company and project names with fallbacks.
    
    Returns:
        tuple: (company_name, project_name)
    """
    company_name = ""
    project_name = ""
    
    # Try multiple patterns for company name
    company_patterns = [
        r'bei (?:der |dem )?([A-Z][a-zA-Z0-9\s&-]+(?:GmbH|AG|KG|UG|SE))',
        r'Unternehmen[:\s]+([A-Z][a-zA-Z0-9\s&-]+(?:GmbH|AG|KG|UG|SE))',
        r'Firma[:\s]+([A-Z][a-zA-Z0-9\s&-]+(?:GmbH|AG|KG|UG|SE))',
        r'([A-Z][a-zA-Z0-9\s&-]+(?:GmbH|AG|KG|UG|SE))',
    ]
    
    for pattern in company_patterns:
        match = re.search(pattern, content)
        if match:
            company_name = match.group(1).strip()
            break
    
    # Try multiple patterns for project name
    project_patterns = [
        r'(?:Projekt|Project)[:\s]+["\']?([A-Z][a-zA-Z0-9\s-]+?)(?:["\']|\.|,|\n)',
        r'(?:Aktuelles Projekt|Current Project)[:\s]+["\']?([A-Z][a-zA-Z0-9\s-]+?)(?:["\']|\.|,|\n)',
        r'"([A-Z][a-zA-Z0-9\s-]+?)"(?:-Projekt|\sproj)',
    ]
    
    for pattern in project_patterns:
        match = re.search(pattern, content)
        if match:
            project_name = match.group(1).strip()
            break
    
    return company_name, project_name


def _calculate_optimal_batch_sizes(topics_count: int, complexity_score: float) -> list:
    """
    Intelligent dynamic batching algorithm.
    
    Industry-standard approach: Adaptive chunk sizing based on:
    - Content volume (number of topics)
    - Complexity (affects solution length)
    - Token budget constraints (8K output tokens per call)
    - Quality preservation (complete solutions)
    
    Strategy:
    - Simple topics (complexity < 4): Batch of 4-5 per pass
    - Medium topics (complexity 4-7): Batch of 3 per pass
    - Complex topics (complexity > 7): Batch of 2 per pass
    - Always keep theory + first batch together
    
    Returns:
        List of batch sizes, e.g., [3, 2, 2] = Pass1: 3 topics, Pass2: 2 topics, Pass3: 2 topics
    """
    if topics_count <= 0:
        return []
    
    # Calculate base batch size from complexity
    if complexity_score < 4.0:
        base_batch = 4  # Simple content: larger batches
    elif complexity_score < 7.0:
        base_batch = 3  # Medium complexity: moderate batches
    else:
        base_batch = 2  # High complexity: smaller batches
    
    # Build batches
    batches = []
    remaining = topics_count
    
    while remaining > 0:
        if remaining <= base_batch:
            batches.append(remaining)
            break
        elif remaining <= base_batch * 1.5:
            # Split remaining evenly
            batch_size = remaining // 2
            batches.append(batch_size)
            batches.append(remaining - batch_size)
            break
        else:
            batches.append(base_batch)
            remaining -= base_batch
    
    return batches


def _estimate_solution_length(topic: Topic, complexity_score: float) -> int:
    """
    Estimate character length for a single problem-solution pair
    
    Based on empirical data from generated scenarios:
    - Simple topic: 1,500-2,000 chars
    - Medium topic: 2,000-3,000 chars
    - Complex topic: 3,000-4,500 chars
    """
    base_length = 2000
    
    # Adjust by complexity
    complexity_multiplier = 0.5 + (complexity_score / 10.0)
    
    # Adjust by topic keywords
    if any(keyword in topic.title.lower() for keyword in ['berechnung', 'kalkulation', 'formel', 'algorithmus']):
        complexity_multiplier *= 1.3  # Math/calculation topics are longer
    
    estimated = int(base_length * complexity_multiplier)
    return min(max(estimated, 1500), 4500)  # Clamp between 1.5K - 4.5K


async def generate_ditele_scenario(
    document_content: str,
    doc_name: str,
    analysis: AnalysisResult
) -> str:
    """
    Generate DiTeLe-compliant educational scenario (ADAPTIVE MULTI-PASS APPROACH)
    
    🚀 INTELLIGENT FEATURES:
    - Dynamic topic selection (no arbitrary limit)
    - Adaptive batching based on complexity
    - Token budget optimization
    - Complete solutions guaranteed
    
    DiTeLe STRUCTURE:
    1. Themenliste (Topic List)
    2. Lernziele (Learning Objectives upfront)
    3. Theoretische Grundlagen (Expanded theory - 700+ words)
    4. Ausgangslage (Realistic starting situation)
    5. Multiple Problem-Solution Pairs (One per major topic)
    6. Lernziel-Checkliste (Learning objectives as checklist)
    
    Returns:
        DiTeLe-formatted scenario content as string
    """
    # 🧠 INTELLIGENT TOPIC SELECTION
    # Select topics based on importance and coverage, not arbitrary limits
    max_topics = min(len(analysis.topics), 7)  # Reasonable upper bound for pedagogy
    topics_list = [t.title for t in analysis.topics[:max_topics]]
    
    # Calculate optimal batch distribution
    batch_sizes = _calculate_optimal_batch_sizes(len(topics_list), analysis.complexity_score)
    total_passes = len(batch_sizes) + 1  # +1 for final CHECKLISTE pass
    
    logger.info(f"Generiere DiTeLe-Szenario...")
    logger.info(f"Abzudeckende Themen: {', '.join(topics_list[:5])}...")
    logger.info(f"Insgesamt {len(topics_list)} Themen -> {len(topics_list)} Problem-Loesungs-Paare")
    
    # Create topic list for prompt
    topics_str = '\n'.join([f"- {t}" for t in topics_list])
    
    # CRITICAL: DiTeLe-COMPLIANT PROMPT (WITHOUT CHECKLISTE - generated separately)
    prompt = f"""Du bist ein erfahrener IT-Ausbilder und Praxisexperte für Fachinformatiker Anwendungsentwicklung.

AUFGABE: Erstelle ein PRAXISNAHES Lernszenario nach dem DiTeLe-Standard

QUELLDOKUMENT: {doc_name}

DOKUMENT-INHALT (Auszug):
{document_content[:3500]}
[... Dokument enthaelt {len(document_content)} Zeichen total ...]

IDENTIFIZIERTE THEMEN (Je Thema = 1 Problem-Loesungs-Paar):
{topics_str}

===============================================================
DiTeLe-STRUKTUR (EXAKT EINHALTEN!)
===============================================================

ABSCHNITT 1: THEMENLISTE
===============================================================

Behandelte Themen:
{topics_str}

===============================================================
ABSCHNITT 2: LERNZIELE (VOR dem Szenario!)
===============================================================

Nach diesem Szenario koennen Sie:
- [Konkrete Faehigkeit 1 - bezogen auf Thema 1]
- [Konkrete Faehigkeit 2 - bezogen auf Thema 2]
- [Konkrete Faehigkeit 3 - bezogen auf Thema 3]
- [Weitere Lernziele fuer jedes Thema - insgesamt {len(topics_list)} Lernziele]

WICHTIG: Formuliere messbare, konkrete Lernziele (z.B. "Kostenarten unterscheiden", nicht "Kostenkalkulation verstehen")

===============================================================
ABSCHNITT 3: THEORETISCHE GRUNDLAGEN
===============================================================

[AUSFUEHRLICHE theoretische Erklaerung - MINDESTENS 700 Woerter!]

Aufbau:
- Absatz 1: Ueberblick ueber die Thematik und Relevanz in der IT-Praxis
- Absatz 2-3: Kernkonzepte der wichtigsten Themen klar erklaert
- Absatz 4-5: Verbindungen zwischen den Themen aufzeigen
- Absatz 6: Praktische Bedeutung und Anwendungsbeispiele

KRITISCH:
- EINFACHE Sprache fuer Auszubildende
- Roter Faden erkennbar
- Keine zu komprimierten Erklaerungen
- Beispiele aus der echten IT-Berufspraxis
- Fachbegriffe werden eingefuehrt und erklaert

===============================================================
ABSCHNITT 4: AUSGANGSLAGE
===============================================================

Du bist Auszubildende/r zum Fachinformatiker fuer Anwendungsentwicklung im [X.] Lehrjahr bei der [Firmenname] GmbH.

Das Unternehmen:
- [15-25] Mitarbeiter
- Standort: [Stadt/Region]
- Geschaeftsfeld: [Entwickelt Softwareloesungen fuer...]
- Spezialisierung: [z.B. Webanwendungen, Mobile Apps, ERP-Systeme...]

Deine Rolle:
- Im Team: [Entwicklerteam, Projektteam, etc.]
- Verantwortlich fuer: [Konkrete Aufgaben]
- Aktuelles Projekt: [Realistisches IT-Projekt das ALLE Themen natuerlich integriert]

Projekthintergrund:
[3-4 Saetze ueber das konkrete Projekt/Problem]
[Welcher Kunde? Welche Anforderungen? Welcher Zeitrahmen?]
[Warum ist dieses Projekt wichtig?]

===============================================================
ABSCHNITT 5: PROBLEME & LOESUNGEN (Je Thema EIN Paar!)
===============================================================

{_generate_problem_solution_template(topics_list)}

===============================================================

KRITISCHE ANFORDERUNGEN:

1. Struktur EXAKT wie oben einhalten (5 Abschnitte - CHECKLISTE kommt spaeter)
2. Fuer JEDES Thema EIN Problem-Loesungs-Paar
3. Theoretische Grundlagen: MINDESTENS 700 Woerter
4. Jede Loesung: Step-by-Step mit WARUM-Erklaerungen
5. JEDE Lösung MUSS VOLLSTÄNDIG sein (ALLE Schritte bis zum Ende - Schritt 1, 2, 3, 4, 5, 6, etc.)
6. KEINE Abbrüche mitten in einer Lösung (z.B. "Schritt 4: Berechne..." und dann STOP)
7. Wenn eine Berechnung 6 Schritte hat, müssen ALLE 6 Schritte da sein!
8. Realistische IT-Szenarien (keine theoretischen Konstrukte)
9. Anfaenger-freundlich: Auszubildende muessen folgen koennen
10. KEINE Markdown-Formatierung (**, ##, ```)
11. KEINE Emojis oder Sonderzeichen
12. KEINE Erwaehnung von "Bot", "KI", "Score", "Qualitaet", "AI"
13. Deutsche Sprache, professionell
14. Trainer-ready: Loesungen detailliert genug fuer Vergleich mit Schuelerarbeit

ABSOLUT VERBOTEN:
- Lösungen die mitten im Rechenvorgang abbrechen
- Unvollständige Step-by-Step Anleitungen
- "..." oder "etc." statt konkrete Schritte auszuführen

ZIELGRUPPE: Auszubildende und Schueler (Anfaenger bis leicht Fortgeschritten)
OUTPUT-ZWECK: Professionelles Lehrmaterial fuer Trainer und Lernende
QUALITAET: Direkt verwendbar ohne Nachbearbeitung

Erstelle jetzt das vollstaendige DiTeLe-Szenario (Abschnitte 1-5):"""

    try:
        # 🚀 INTELLIGENT ADAPTIVE MULTI-PASS GENERATION
        logger.info(f"   🧠 Batching-Strategie: {batch_sizes} (Komplexität: {analysis.complexity_score:.1f}/10)")
        logger.info(f"   📊 Gesamt: {total_passes} Passe ({len(batch_sizes)} für Probleme + 1 für CHECKLISTE)")
        
        main_content = ""
        processed_topics = 0
        
        # Context variables to maintain consistency across batches
        company_name = ""
        project_name = ""
        
        # Generate each batch dynamically
        for pass_num, batch_size in enumerate(batch_sizes, 1):
            batch_topics = topics_list[processed_topics:processed_topics + batch_size]
            is_first_pass = (pass_num == 1)
            
            # Calculate problem numbers for this batch
            start_problem_num = processed_topics + 1
            end_problem_num = processed_topics + batch_size
            
            logger.info(f"   [PASS {pass_num}/{total_passes}] Generiere Problem {start_problem_num}-{end_problem_num} ({len(batch_topics)} Paare){' (+ Hauptinhalt)' if is_first_pass else ''}...")
            
            if is_first_pass:
                # First pass includes theory + starting situation + first batch
                batch_prompt = prompt.replace(
                    f"{_generate_problem_solution_template(topics_list)}",
                    f"{_generate_problem_solution_template(batch_topics, start_num=1)}"
                    + (f"\n\n[Weitere {len(topics_list) - batch_size} Problem-Loesungs-Paare folgen in weiteren Pässen]" 
                       if len(topics_list) > batch_size else "")
                )
                
                batch_content = await intelligent_gemini.generate_from_prompt(
                    prompt=batch_prompt,
                    content_type=f"ditele_pass_{pass_num}_main",
                    timeout=300,
                    max_retries=3
                )
                
                if not batch_content or len(batch_content) < 4000:
                    logger.warning(f"   Zu wenig Inhalt ({len(batch_content) if batch_content else 0} Zeichen), Retry...")
                    retry_prompt = batch_prompt.replace(document_content[:3500], document_content[:2500])
                    batch_content = await intelligent_gemini.generate_from_prompt(
                        prompt=retry_prompt,
                        content_type=f"ditele_pass_{pass_num}_retry",
                        timeout=300
                    )
                
                if not batch_content or len(batch_content) < 3000:
                    logger.error("   Hauptinhalt-Generierung fehlgeschlagen")
                    return ""
                
                # MINIMAL CLEANUP: Remove forbidden terms
                batch_content = re.sub(r'\b[Bb]ot\b', 'System', batch_content)
                batch_content = re.sub(r'\b[Kk][Ii]\b', 'System', batch_content)
                batch_content = re.sub(r'\b[Aa][Ii]\b', 'System', batch_content)
                
                main_content = batch_content
                
                # Extract context for subsequent batches (simple, single-pattern approach)
                company_match = re.search(r'bei (?:der |dem )?([A-Z][a-zA-Z0-9\s&-]+(?:GmbH|AG|KG|UG|SE))', batch_content)
                if company_match:
                    company_name = company_match.group(1).strip()
                    logger.info(f"   📝 Firmenname: {company_name}")
                
                project_match = re.search(r'(?:Projekt|Project)[:\s]+["\']?([A-Z][a-zA-Z0-9\s-]+?)(?:["\']|\.|,|\n)', batch_content)
                if project_match:
                    project_name = project_match.group(1).strip()
                    logger.info(f"   📝 Projektname: {project_name}")
                
            else:
                # Subsequent passes: only problem-solution pairs with CONTEXT
                completed_topics_str = '\n'.join([f"- {t}" for t in topics_list[:processed_topics]])
                batch_topics_str = '\n'.join([f"- {t}" for t in batch_topics])
                
                # Build context string
                context_info = f"""WICHTIGER KONTEXT (aus vorherigen Abschnitten):
- Unternehmen: {company_name if company_name else '[Name aus vorherigem Abschnitt übernehmen]'}
- Projekt: {project_name if project_name else '[Name aus vorherigem Abschnitt übernehmen]'}
- Bereits behandelt: Problem 1-{processed_topics}"""
                
                batch_prompt = f"""Du bist ein erfahrener IT-Ausbilder und Praxisexperte für Fachinformatiker Anwendungsentwicklung.

AUFGABE: Erstelle die NÄCHSTEN Problem-Lösungs-Paare für das DiTeLe-Szenario

{context_info}

BEREITS BEARBEITETE THEMEN ({processed_topics} Stück):
{completed_topics_str}

JETZT ZU BEARBEITEN ({len(batch_topics)} Themen - Problem {start_problem_num} bis {end_problem_num}):
{batch_topics_str}

DOKUMENT-INHALT (Auszug):
{document_content[:3000]}

WICHTIG: 
- Verwende DIESELBEN Namen (Firma, Projekt) wie in den vorherigen Abschnitten
- Nummerierung: Starte bei PROBLEM {start_problem_num} und ende bei PROBLEM {end_problem_num}
- Halte den gleichen Stil und die gleiche Qualität wie in den bisherigen Lösungen bei

Erstelle für jedes der {len(batch_topics)} Themen EIN vollständiges Problem-Lösungs-Paar:

{_generate_problem_solution_template(batch_topics, start_num=start_problem_num)}

KRITISCH - VOLLSTÄNDIGKEIT:
- JEDE Lösung muss VOLLSTÄNDIG sein (ALLE Schritte bis zum Ende!)
- Wenn ein Problem 6 Berechnungsschritte benötigt, müssen ALLE 6 da sein
- Schritt 4, 5, 6, etc. MÜSSEN vorhanden sein bis die Lösung komplett ist
- NIEMALS mitten in einem Schritt aufhören (z.B. "Schritt 4: Berechne..." und dann STOP)
- KEINE Abbrüche mitten in der Lösung
- Beispiel für FALSCH: "Schritt 3: Berechne die Basis... Herstellkosten = 1.872.000 EUR. Schritt 4: Berechne den Verwaltungsgemeinkost" [STOP] <- FALSCH!
- Beispiel für RICHTIG: "Schritt 3... Schritt 4: Berechne VwGK-Satz = ... = 8,01%. Schritt 5: Berechne VtGK-Satz = ... = 2,67%. Ergebnis: [Alle 4 Sätze vollständig]"

KRITISCH - KONSISTENZ:
- Verwende EXAKT die gleichen Namen wie vorher (Firma: {company_name if company_name else '[aus Pass 1]'}, Projekt: {project_name if project_name else '[aus Pass 1]'})
- Nummeriere die Probleme fortlaufend: PROBLEM {start_problem_num}, PROBLEM {start_problem_num + 1}, etc.
- NICHT bei 1 neu anfangen!
- Step-by-Step mit WARUM-Erklaerungen
- KEINE Markdown-Formatierung (**, ##, ```)
- KEINE Emojis oder Sonderzeichen
- KEINE Erwähnung von "Bot", "KI", "Score", "AI"
- Deutsche Sprache, professionell

Erstelle jetzt die vollständigen Problem-Lösungs-Paare (PROBLEM {start_problem_num} bis PROBLEM {end_problem_num}):"""
                
                batch_content = await intelligent_gemini.generate_from_prompt(
                    prompt=batch_prompt,
                    content_type=f"ditele_pass_{pass_num}_batch",
                    timeout=300,
                    max_retries=2
                )
                
                if batch_content and len(batch_content) > 500:
                    # CLEANUP: Remove forbidden terms AND renumber problems
                    batch_content = re.sub(r'\b[Bb]ot\b', 'System', batch_content)
                    batch_content = re.sub(r'\b[Kk][Ii]\b', 'System', batch_content)
                    batch_content = re.sub(r'\b[Aa][Ii]\b', 'System', batch_content)
                    
                    # Renumber problems to maintain sequential order
                    batch_content = _cleanup_batch_content(batch_content, start_problem_num, end_problem_num)
                    
                    main_content += "\n\n" + batch_content
                    logger.info(f"   [PASS {pass_num}/{total_passes}] Erfolgreich: {len(batch_content):,} Zeichen hinzugefügt")
                else:
                    logger.warning(f"   [PASS {pass_num}/{total_passes}] Batch zu kurz oder fehlgeschlagen, fahre trotzdem fort")
            
            processed_topics += batch_size
            logger.info(f"   ✅ Fortschritt: {processed_topics}/{len(topics_list)} Themen abgeschlossen")
        
        logger.info(f"   📦 Gesamtinhalt nach allen Batches: {len(main_content):,} Zeichen")
        
        # FINAL PASS: Generate CHECKLISTE (Section 6)
        logger.info(f"   [PASS {total_passes}/{total_passes}] Generiere Lernziel-Checkliste...")
        
        # Extract Lernziele from main content
        lernziele_match = re.search(r'LERNZIELE.*?Nach diesem Szenario.*?:(.*?)(?:THEORETISCHE GRUNDLAGEN|Theoretische Grundlagen)', 
                                   main_content, re.DOTALL | re.IGNORECASE)
        lernziele_text = lernziele_match.group(1).strip() if lernziele_match else ""
        
        checkliste_prompt = f"""Erstelle eine Lernziel-Checkliste basierend auf folgenden Lernzielen.

LERNZIELE:
{lernziele_text if lernziele_text else topics_str}

AUFGABE:
Formuliere jedes Lernziel als eine Frage mit "Koennen Sie...?" oder "Sind Sie in der Lage...?"

FORMAT:
===============================================================
ABSCHNITT 6: LERNZIEL-CHECKLISTE
===============================================================

Koennen Sie jetzt...?

[  ] [Lernziel 1 als Frage formuliert]
[  ] [Lernziel 2 als Frage formuliert]
[  ] [Lernziel 3 als Frage formuliert]
[  ] [Alle weiteren Lernziele als Fragen]

WICHTIG:
- Verwende [  ] statt Emojis
- Jede Frage beginnt mit Grossbuchstaben
- Fragen sind praezise und messbar
- KEINE Emojis oder Sonderzeichen
- Deutsche Sprache

Erstelle jetzt die Checkliste:"""

        checkliste = await intelligent_gemini.generate_from_prompt(
            prompt=checkliste_prompt,
            content_type="ditele_checkliste",
            timeout=120,
            max_retries=2
        )
        
        if checkliste and len(checkliste) > 100:
            logger.info(f"   [PASS {total_passes}/{total_passes}] Checkliste erfolgreich: {len(checkliste)} Zeichen")
            # Combine main content + checkliste
            full_content = main_content + "\n\n" + checkliste
        else:
            logger.warning(f"   [PASS {total_passes}/{total_passes}] Checkliste-Generierung fehlgeschlagen, verwende Fallback")
            # Fallback: Create simple checkliste from topics
            fallback_checkliste = "\n\n===============================================================\n"
            fallback_checkliste += "ABSCHNITT 6: LERNZIEL-CHECKLISTE\n"
            fallback_checkliste += "===============================================================\n\n"
            fallback_checkliste += "Koennen Sie jetzt...?\n\n"
            for i, topic in enumerate(topics_list, 1):
                fallback_checkliste += f"[  ] das Thema '{topic}' in der Praxis anwenden?\n"
            full_content = main_content + fallback_checkliste
        
        logger.info(f"   Gesamtinhalt: {len(full_content):,} Zeichen")
        return full_content
    
    except Exception as e:
        logger.error(f"   Fehler bei Generierung: {e}")
        import traceback
        traceback.print_exc()
        return ""


def validate_ditele_structure(
    scenario_content: str,
    analysis: AnalysisResult
) -> Dict[str, Any]:
    """
    🔍 Validate DiTeLe structure completeness and quality
    
    Checks for all 6 required DiTeLe sections and content quality
    
    Returns:
        Dict with validation results, issues, and quality metrics
    """
    logger.info("🔍 Führe DiTeLe-Strukturprüfung durch...")
    
    issues = []
    warnings = []
    
    # Check content length (DiTeLe scenarios are longer)
    content_length = len(scenario_content)
    if content_length < 5000:
        issues.append(f"Zu wenig Inhalt: {content_length} Zeichen (Minimum: 5000 für DiTeLe)")
    elif content_length < 6000:
        warnings.append(f"Inhalt könnte ausführlicher sein: {content_length} Zeichen (Ziel: 6000+)")
    else:
        logger.info(f"   ✅ Länge ausreichend: {content_length:,} Zeichen")
    
    # Check for all 6 DiTeLe sections
    required_sections = [
        'THEMENLISTE',
        'LERNZIELE',
        'THEORETISCHE GRUNDLAGEN',
        'AUSGANGSLAGE',
        'PROBLEM',
        'LÖSUNG',
        'CHECKLISTE'
    ]
    
    missing_sections = []
    for section in required_sections:
        if section.upper() not in scenario_content.upper():
            missing_sections.append(section)
    
    if missing_sections:
        issues.append(f"Fehlende DiTeLe-Abschnitte: {', '.join(missing_sections)}")
    else:
        logger.info(f"   ✅ Alle DiTeLe-Abschnitte vorhanden")
    
    # Count problem-solution pairs
    problem_count = scenario_content.upper().count("PROBLEM")
    solution_count = scenario_content.upper().count("LÖSUNG")
    expected_pairs = min(len(analysis.topics), 7)
    
    if problem_count < expected_pairs:
        warnings.append(f"Nur {problem_count}/{expected_pairs} Probleme gefunden")
    if solution_count < expected_pairs:
        warnings.append(f"Nur {solution_count}/{expected_pairs} Lösungen gefunden")
    else:
        logger.info(f"   ✅ {problem_count} Problem-Lösungs-Paare gefunden")
    
    # Check theory section length (should be 800+ words)
    theory_indicators = ['theoretische', 'grundlagen', 'konzept', 'prinzip']
    has_theory = any(ind in scenario_content.lower() for ind in theory_indicators)
    if not has_theory:
        issues.append("Theoretische Grundlagen fehlen oder zu kurz")
    else:
        logger.info(f"   ✅ Theoretische Grundlagen vorhanden")
    
    # Check for forbidden terms
    forbidden_terms = ['bot', 'ki-generiert', 'quality score', 'qualitätsscore', 'ai-generated']
    for term in forbidden_terms:
        if term in scenario_content.lower():
            issues.append(f"Unerlaubte Erwähnung: '{term}'")
    
    # Check for markdown symbols (should be minimal)
    markdown_symbols = ['**', '##', '###', '```']
    markdown_count = sum(scenario_content.count(sym) for sym in markdown_symbols)
    if markdown_count > 10:
        warnings.append(f"Viele Markdown-Symbole gefunden ({markdown_count}x)")
    
    # Calculate quality score
    quality_score = 100
    quality_score -= len(issues) * 20
    quality_score -= len(warnings) * 10
    quality_score = max(0, min(100, quality_score))
    
    validation_result = {
        'quality_score': quality_score,
        'problem_solution_pairs': problem_count,
        'expected_pairs': expected_pairs,
        'content_length': content_length,
        'issues': issues,
        'warnings': warnings,
        'passed': len(issues) == 0,
        'grade': 'EXCELLENT' if quality_score >= 90 else 'GOOD' if quality_score >= 75 else 'ACCEPTABLE' if quality_score >= 60 else 'NEEDS_IMPROVEMENT'
    }
    
    # Log results
    if validation_result['passed']:
        logger.info(f"   ✅ DiTeLe-Strukturprüfung BESTANDEN")
        logger.info(f"   🎯 Interner Score: {quality_score}/100 ({validation_result['grade']})")
        logger.info(f"   💡 Problem-Lösungs-Paare: {problem_count}/{expected_pairs}")
    else:
        logger.warning(f"   ⚠️ DiTeLe-Strukturprüfung mit Problemen")
        logger.warning(f"   🎯 Interner Score: {quality_score}/100 ({validation_result['grade']})")
        for issue in issues:
            logger.warning(f"      ❌ {issue}")
    
    for warning in warnings:
        logger.info(f"      ⚠️ {warning}")
    
    return validation_result


def create_ditele_word_document(
    scenario_content: str,
    filename: str,
    analysis: AnalysisResult,
    validation: Dict[str, Any]
) -> bytes:
    """
    📄 Create professional Word document with DiTeLe structure
    
    Structure:
    - Title page
    - Themenliste (Topic List)
    - Lernziele (Learning Objectives)
    - Theoretische Grundlagen (Theoretical Foundation)
    - Ausgangslage (Starting Situation)
    - Problem-Lösungs-Paare (Multiple Problem-Solution Pairs)
    - Lernziel-Checkliste (Learning Checklist)
    - Footer with metadata
    """
    logger.info("Erstelle DiTeLe Word-Dokument...")
    
    doc = Document()
    
    # ========== EMOJI REMOVAL FUNCTION ==========
    def remove_emojis(text: str) -> str:
        """Remove all emojis and special unicode characters from text"""
        import re
        # Pattern for emojis and various unicode symbols
        emoji_pattern = re.compile("["
            u"\U0001F600-\U0001F64F"  # emoticons
            u"\U0001F300-\U0001F5FF"  # symbols & pictographs
            u"\U0001F680-\U0001F6FF"  # transport & map symbols
            u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
            u"\U00002702-\U000027B0"
            u"\U000024C2-\U0001F251"
            u"\U0001F900-\U0001F9FF"  # Supplemental Symbols and Pictographs
            u"\U0001FA00-\U0001FA6F"  # Chess Symbols
            u"\U00002600-\U000026FF"  # Miscellaneous Symbols
            u"\U00002700-\U000027BF"  # Dingbats
            "]+", flags=re.UNICODE)
        text = emoji_pattern.sub('', text)
        # Remove checkboxes and special symbols
        text = text.replace('☐', '[  ]').replace('✓', '[x]').replace('✅', '[x]')
        text = text.replace('❌', '[!]').replace('⚠️', '[!]').replace('📚', '').replace('🎯', '')
        text = text.replace('═', '=').replace('─', '-').replace('│', '|')
        return text.strip()
    
    # ========== PROFESSIONAL STYLING ==========
    # Normal text style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(11)
    
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15
    
    # Heading 1 style
    heading1_style = doc.styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Arial Narrow'
    heading1_font.size = Pt(18)
    heading1_font.bold = True
    heading1_font.color.rgb = RGBColor(0, 51, 153)
    heading1_style.paragraph_format.space_before = Pt(12)
    heading1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style
    heading2_style = doc.styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = 'Arial Narrow'
    heading2_font.size = Pt(14)
    heading2_font.bold = True
    heading2_font.color.rgb = RGBColor(0, 102, 204)
    heading2_style.paragraph_format.space_before = Pt(10)
    heading2_style.paragraph_format.space_after = Pt(4)
    
    # Heading 3 style
    heading3_style = doc.styles['Heading 3']
    heading3_font = heading3_style.font
    heading3_font.name = 'Arial Narrow'
    heading3_font.size = Pt(12)
    heading3_font.bold = True
    heading3_font.color.rgb = RGBColor(0, 120, 215)
    heading3_style.paragraph_format.space_before = Pt(8)
    heading3_style.paragraph_format.space_after = Pt(3)
    
    # ========== TITLE PAGE ==========
    # Extract source document number from filename (e.g., "01_03_FR-660.docx" -> "01_03_FR-660")
    source_number = filename.replace('.docx', '').replace('.pdf', '').replace('.txt', '')
    
    title = doc.add_heading(f'Enabler Scenario {source_number}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = 'Arial Narrow'
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.color.rgb = RGBColor(0, 51, 153)
    title.runs[0].font.bold = True
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph(f'Basierend auf: {filename}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.name = 'Arial Narrow'
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    subtitle.runs[0].italic = True
    subtitle.paragraph_format.space_before = Pt(12)
    subtitle.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    standard_para = doc.add_paragraph('DiTeLe-Standard')
    standard_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    standard_para.runs[0].font.name = 'Arial Narrow'
    standard_para.runs[0].font.size = Pt(13)
    standard_para.runs[0].font.italic = True
    standard_para.runs[0].font.color.rgb = RGBColor(70, 130, 180)
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph(f'Erstellt am: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.name = 'Arial Narrow'
    date_para.runs[0].font.size = Pt(12)
    date_para.runs[0].font.italic = True
    date_para.paragraph_format.space_before = Pt(6)
    
    doc.add_page_break()
    
    # ========== PARSE AND FORMAT DITELE CONTENT ==========
    _parse_and_format_ditele_content(doc, scenario_content, remove_emojis)
    
    # ========== FOOTER ==========
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f'FIAE DiTeLe-Szenario | Generiert am {datetime.now().strftime("%d.%m.%Y %H:%M")} | Seite '
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.name = 'Arial Narrow'
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Save as bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    doc_size_kb = len(buffer.getvalue()) / 1024
    logger.info(f"✅ Word-Dokument erstellt: {doc_size_kb:.1f} KB")
    
    return buffer.getvalue()


def _parse_and_format_ditele_content(doc: Document, content: str, remove_emoji_func):
    """
    Parse AI-generated DiTeLe content and format it properly in Word document
    
    Handles:
    - Section headings (Themenliste, Lernziele, etc.)
    - Problem-Solution pairs with proper formatting
    - Bullet lists
    - Step-by-step instructions
    - Checklists
    - Emoji removal
    """
    # Remove all emojis from content first
    content = remove_emoji_func(content)
    
    current_section = None
    in_solution_steps = False
    
    for line in content.split('\n'):
        line = line.strip()
        if not line or line == '=' * len(line) or line == '-' * len(line):
            continue
        
        # Detect section headings
        line_upper = line.upper()
        
        # Main sections (Heading 1)
        if any(section in line_upper for section in ['THEMENLISTE', 'LERNZIELE', 'THEORETISCHE GRUNDLAGEN', 'AUSGANGSLAGE', 'CHECKLISTE']):
            doc.add_paragraph()
            if 'THEMENLISTE' in line_upper:
                doc.add_heading('Behandelte Themen', level=1)
                current_section = 'themenliste'
            elif 'LERNZIELE' in line_upper:
                doc.add_heading('Lernziele', level=1)
                current_section = 'lernziele'
            elif 'THEORETISCHE GRUNDLAGEN' in line_upper:
                doc.add_heading('Theoretische Grundlagen', level=1)
                current_section = 'theorie'
            elif 'AUSGANGSLAGE' in line_upper:
                doc.add_heading('Ausgangslage', level=1)
                current_section = 'ausgangslage'
            elif 'CHECKLISTE' in line_upper:
                doc.add_heading('Lernziel-Checkliste', level=1)
                current_section = 'checkliste'
            continue
        
        # Problem headings (Heading 2)
        if line_upper.startswith('PROBLEM'):
            doc.add_paragraph()
            heading = doc.add_heading(line, level=2)
            current_section = 'problem'
            continue
        
        # Solution headings (Heading 3)
        if line_upper.startswith('LÖSUNG'):
            doc.add_paragraph()
            heading = doc.add_heading(line, level=3)
            current_section = 'lösung'
            in_solution_steps = True
            continue
        
        # Sub-headings within sections
        if line_upper.startswith(('SITUATION:', 'AUFGABE:', 'RANDBEDINGUNGEN:', 'ERWARTETE ERGEBNISSE:', 
                                  'ERGEBNIS:', 'ERKLÄRUNG:', 'ALTERNATIVE ANSÄTZE:', 'HÄUFIGE FEHLER:')):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.name = 'Arial Narrow'
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            continue
        
        # Step-by-step instructions (Schritt 1, Schritt 2, etc.)
        if re.match(r'^Schritt \d+:', line, re.IGNORECASE):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.name = 'Arial Narrow'
            run.font.color.rgb = RGBColor(0, 102, 204)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            continue
        
        # Bullet points or checklists
        if line.startswith(('-', '•', '☐', '✓', '✔')):
            bullet_para = doc.add_paragraph(line[1:].strip(), style='List Bullet')
            for run in bullet_para.runs:
                run.font.name = 'Arial Narrow'
            continue
        
        # Normal paragraphs
        if len(line) > 2:
            normal_para = doc.add_paragraph(line)
            for run in normal_para.runs:
                run.font.name = 'Arial Narrow'


async def process_document_ditele(doc: Dict[str, Any], index: int, total: int) -> Dict[str, Any]:
    """
    📄 Process single document with DiTeLe structure
    
    Workflow:
    1. Extract document content
    2. Analyze document (AI-powered, identify topics)
    3. Generate DiTeLe-compliant scenario
    4. Validate structure completeness
    5. Create professional Word document
    6. Upload to Google Drive
    
    Returns:
        Dict with status, file info, and metrics
    """
    doc_id = doc['id']
    doc_name = doc['name']
    
    logger.info(f"\n{'='*80}")
    logger.info(f"📄 [{index}/{total}] VERARBEITE: {doc_name}")
    logger.info(f"{'='*80}")
    
    start_time = datetime.now()
    
    try:
        # STEP 1: Extract content
        logger.info("\n[1/5] 📥 Extrahiere Dokumentinhalt...")
        document_content = google_drive_service.get_file_content(doc_id)
        
        if not document_content or len(document_content) < 100:
            logger.error("❌ Zu wenig Inhalt extrahiert")
            return {'status': 'failed', 'error': 'insufficient_content'}
        
        word_count = len(document_content.split())
        logger.info(f"   ✅ {len(document_content):,} Zeichen extrahiert ({word_count:,} Wörter)")
        
        # STEP 2: Analyze document
        logger.info("\n[2/5] 🧠 Analysiere Dokument mit KI...")
        analysis = await analyze_document_for_ditele(document_content, doc_name)
        
        # STEP 3: Generate DiTeLe scenario
        logger.info(f"\n[3/5] 🤖 Generiere DiTeLe-Szenario...")
        scenario_content = await generate_ditele_scenario(
            document_content,
            doc_name,
            analysis
        )
        
        if not scenario_content or len(scenario_content) < 4000:
            logger.error("❌ DiTeLe-Generierung fehlgeschlagen oder zu wenig Inhalt")
            return {'status': 'failed', 'error': 'generation_failed'}
        
        # STEP 4: Validate DiTeLe structure
        logger.info("\n[4/5] 🔍 Führe DiTeLe-Strukturprüfung durch...")
        validation = validate_ditele_structure(scenario_content, analysis)
        
        if not validation['passed']:
            logger.warning("⚠️ DiTeLe-Strukturprüfung mit Problemen, fahre aber fort...")
        
        # STEP 5: Create Word document and upload
        logger.info("\n[5/5] 💾 Erstelle und lade Word-Dokument hoch...")
        
        doc_bytes = create_ditele_word_document(
            scenario_content,
            doc_name,
            analysis,
            validation
        )
        
        # Upload to Google Drive
        base_name = doc_name.rsplit('.', 1)[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"Enabler_Scenario_{base_name}_{timestamp}.docx"
        
        review_folder_id = settings.google_drive_review_folder_id
        
        file_metadata = {
            'name': output_filename,
            'parents': [review_folder_id],
            'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        
        if not google_drive_service or not google_drive_service.service:
            raise Exception("Google Drive service not initialized")
        
        media = MediaIoBaseUpload(
            io.BytesIO(doc_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            resumable=True
        )
        
        file = google_drive_service.service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id,name,webViewLink'
        ).execute()
        
        elapsed = (datetime.now() - start_time).total_seconds()
        
        logger.info(f"\n✅ ERFOLG! Abgeschlossen in {elapsed:.1f}s")
        logger.info(f"📁 Datei: {output_filename}")
        logger.info(f"🔗 Link: {file.get('webViewLink', 'N/A')}")
        logger.info(f"📄 Format: DiTeLe-Standard")
        logger.info(f"💡 Problem-Lösungs-Paare: {validation['problem_solution_pairs']}")
        logger.info(f"✨ Professional, anfängerfreundlich, trainer-ready")
        
        return {
            'status': 'completed',
            'doc_name': doc_name,
            'file_id': file['id'],
            'file_name': output_filename,
            'web_view_link': file.get('webViewLink'),
            'problem_solution_pairs': validation['problem_solution_pairs'],
            'content_length': validation['content_length'],
            'processing_time': elapsed
        }
        
    except Exception as e:
        logger.error(f"❌ Fehler bei Verarbeitung von {doc_name}: {e}")
        import traceback
        traceback.print_exc()
        return {
            'status': 'failed',
            'doc_name': doc_name,
            'error': str(e)
        }


async def main():
    """
    Main workflow for DiTeLe scenario generation
    
    Supports TEST_MODE for quality review before full processing
    """
    
    logger.info("="*80)
    logger.info("🚀 PHASE 1: DiTeLe STANDARD SCENARIO GENERATION")
    logger.info("="*80)
    
    if TEST_MODE:
        logger.warning(f"⚠️  TEST MODE AKTIV - Verarbeite nur {MAX_TEST_DOCUMENTS} Dokument(e)")
        logger.warning(f"⚠️  Setze TEST_MODE = False für vollständige Verarbeitung")
    
    logger.info("🎯 Struktur: DiTeLe-Standard (6 Abschnitte)")
    logger.info("📄 Output: Professional Word-Dokumente")
    logger.info("💡 Fokus: Multiple Problem-Lösungs-Paare pro Dokument")
    logger.info("✅ Qualität: Anfängerfreundlich, trainer-ready, detailliert")
    logger.info("="*80 + "\n")
    
    try:
        # Discover documents
        source_folder_id = settings.google_drive_content_source_folder_id
        logger.info(f"🔍 Suche Dokumente in Quellordner...")
        
        all_documents = google_drive_service.list_files_in_folder(source_folder_id)
        
        # Check if documents are in root or subfolders
        direct_documents = [
            doc for doc in all_documents
            if doc['name'].lower().endswith(('.docx', '.pdf', '.txt'))
            and not doc['name'].startswith('~')
        ]
        
        documents = []
        
        if direct_documents:
            documents = direct_documents
            logger.info(f"📚 {len(documents)} Dokumente im Hauptordner gefunden")
        else:
            # Search in subfolders
            logger.info(f"📁 Keine Dokumente im Hauptordner, suche in Unterordnern...")
            folders = [f for f in all_documents if f.get('mimeType') == 'application/vnd.google-apps.folder']
            logger.info(f"   Gefunden: {len(folders)} Unterordner")
            
            for folder in folders:
                try:
                    folder_files = google_drive_service.list_files_in_folder(folder['id'])
                    folder_docs = [
                        doc for doc in folder_files
                        if doc['name'].lower().endswith(('.docx', '.pdf', '.txt'))
                        and not doc['name'].startswith('~')
                    ]
                    documents.extend(folder_docs)
                except Exception as e:
                    logger.warning(f"   ⚠️ Fehler beim Lesen von Ordner {folder['name']}: {e}")
            
            logger.info(f"📚 {len(documents)} Dokumente in Unterordnern gefunden")
        
        if not documents:
            logger.warning("⚠️ Keine Dokumente im Quellordner gefunden!")
            return
        
        # Apply TEST_MODE limit
        if TEST_MODE:
            original_count = len(documents)
            documents = documents[:MAX_TEST_DOCUMENTS]
            logger.warning(f"🧪 TEST MODE: Verarbeite {len(documents)} von {original_count} Dokumenten")
        
        # Process documents
        results = []
        success_count = 0
        
        for i, doc in enumerate(documents, 1):
            result = await process_document_ditele(doc, i, len(documents))
            results.append(result)
            
            if result['status'] == 'completed':
                success_count += 1
            
            # Pause between documents (rate limiting)
            if i < len(documents):
                logger.info(f"\n⏸️  Pause 15s vor nächstem Dokument...\n")
                await asyncio.sleep(15)
        
        # Final statistics
        logger.info(f"\n{'='*80}")
        logger.info(f"🎉 VERARBEITUNG ABGESCHLOSSEN")
        logger.info(f"{'='*80}")
        logger.info(f"✅ Erfolgreich: {success_count}/{len(documents)}")
        logger.info(f"❌ Fehlgeschlagen: {len(documents) - success_count}/{len(documents)}")
        success_rate = (success_count / len(documents) * 100) if documents else 0
        logger.info(f"📊 Erfolgsquote: {success_rate:.1f}%")
        
        if TEST_MODE:
            logger.warning(f"\n⚠️  TEST MODE war aktiv!")
            logger.warning(f"⚠️  Prüfe die {success_count} generierten Dokumente")
            logger.warning(f"⚠️  Wenn Qualität gut → Setze TEST_MODE = False und starte erneut")
        
        if success_count > 0:
            successful_results = [r for r in results if r['status'] == 'completed']
            
            avg_pairs = sum(r.get('problem_solution_pairs', 0) for r in successful_results) / success_count
            avg_length = sum(r.get('content_length', 0) for r in successful_results) / success_count
            avg_time = sum(r['processing_time'] for r in successful_results) / success_count
            
            logger.info(f"\n📈 Durchschnittswerte:")
            logger.info(f"   📄 Format: DiTeLe-Standard")
            logger.info(f"   💡 Problem-Lösungs-Paare: {avg_pairs:.1f}")
            logger.info(f"   📏 Durchschnittliche Länge: {avg_length:.0f} Zeichen")
            logger.info(f"   ⏱️  Verarbeitungszeit: {avg_time:.1f}s")
            logger.info(f"   ✨ Professional, strukturiert, anfängerfreundlich")
        
        logger.info(f"{'='*80}\n")
        
    except Exception as e:
        logger.error(f"❌ Fataler Fehler: {e}")
        import traceback
        traceback.print_exc()
        raise


if __name__ == "__main__":
    asyncio.run(main())
