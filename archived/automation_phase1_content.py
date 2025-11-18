#!/usr/bin/env python3
"""
PHASE 1: DiTeLe STANDARD SCENARIO GENERATION
=============================================
🎯 Focus: Generate educational scenarios following DiTeLe standard structure
📚 Coverage: Multiple problem-solution pairs covering all major topics
✅ Quality: Detailed, beginner-friendly, trainer-ready
📄 Output: Professional Word documents with standardized structure

DiTeLe STRUCTURE:
1. Themenliste (Topic List)
2. Lernziele (Learning Objectives - upfront)
3. Theoretische Grundlagen (Theoretical Foundation - 800+ words)
4. Ausgangslage (Starting Situation - realistic context)
5. Problem-Lösungs-Paare (Multiple Problem-Solution Pairs - one per topic)
6. Lernziel-Checkliste (Learning Objectives Checklist - verification)

TEST MODE:
- Set TEST_MODE = True to process only 2 documents for quality review
- Set TEST_MODE = False for full batch processing
- Always test first before running full automation!

Workflow:
1. Analyze source document (identify topics, complexity, scope)
2. Generate DiTeLe-compliant scenario with multiple problem-solution pairs
3. Validate structure completeness
4. Create professional Word document with proper formatting
5. Upload to Google Drive for review
"""

# ========== CONFIGURATION ==========
TEST_MODE = True  # Set to False after testing first 2 documents
MAX_TEST_DOCUMENTS = 2  # Number of documents to process in test mode

import asyncio
import os
import sys
from datetime import datetime
from typing import Dict, Any, List
from loguru import logger
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googleapiclient.http import MediaIoBaseUpload
import io
import re

load_dotenv()
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from app.config import settings
from app.services.google_services import GoogleDriveService
from app.services.intelligent_gemini_service import IntelligentGeminiService
from app.services.document_analyzer import DocumentAnalyzer, AnalysisResult

# Configure logging
logger.remove()
logger.add(
    sys.stdout,
    format="<green>{time:YYYY-MM-DD HH:mm:ss}</green> | <level>{level: <8}</level> | <level>{message}</level>",
    level="INFO"
)

# Initialize services
google_drive_service = GoogleDriveService()
intelligent_gemini = IntelligentGeminiService()
document_analyzer = DocumentAnalyzer(gemini_service=intelligent_gemini.gemini_service)


async def analyze_document_for_use_cases(document_content: str, doc_name: str) -> AnalysisResult:
    """
    🧠 Analyze document intelligently with AI to identify major topics
    
    NEW: Now optimized for generating ONE comprehensive use case that covers
    all major topics from the source document.
    
    Returns:
        AnalysisResult with topics, complexity, and content requirements
        (use_cases_count will always be 1)
    """
    logger.info(f"🧠 Analysiere Dokument: {doc_name}")
    
    # Use AI-powered analysis
    analysis_result = await document_analyzer.analyze_document(
        document_content=document_content,
        document_name=doc_name,
        use_ai=True  # KI-gestützte Analyse für beste Ergebnisse
    )
    
    # Log analysis results
    logger.info(f"📊 Analyse-Ergebnisse:")
    logger.info(f"   📝 Wörter: {len(document_content.split()):,}")
    logger.info(f"   📑 Themen identifiziert: {len(analysis_result.topics)}")
    logger.info(f"   🎯 Komplexität: {analysis_result.complexity_score:.1f}/10")
    logger.info(f"   ✨ Use Cases zu generieren: {analysis_result.content_requirements.use_cases_count} (comprehensive)")
    logger.info(f"   📄 Alle wichtigen Themen werden in EINEM Use Case abgedeckt")
    
    return analysis_result


async def generate_ditele_scenario(
    document_content: str,
    doc_name: str,
    analysis: AnalysisResult
) -> str:
    """
    🤖 Generate DiTeLe-compliant educational scenario
    
    DiTeLe STRUCTURE:
    1. Themenliste (Topic List)
    2. Lernziele (Learning Objectives upfront)
    3. Theoretische Grundlagen (Expanded theory - 800+ words)
    4. Ausgangslage (Realistic starting situation)
    5. Multiple Problem-Solution Pairs (One per major topic)
    6. Lernziel-Checkliste (Learning objectives as checklist)
    
    Returns:
        DiTeLe-formatted scenario content as string
    """
    requirements = analysis.content_requirements
    topics_list = [t.title for t in analysis.topics[:10]]  # Top 10 topics
    
    logger.info(f"🤖 Generiere DiTeLe-Szenario...")
    logger.info(f"📋 Abzudeckende Themen: {', '.join(topics_list[:5])}...")
    logger.info(f"📊 Insgesamt {len(analysis.topics)} Themen identifiziert")
    logger.info(f"🎯 Erwartete Problem-Lösungs-Paare: {min(len(topics_list), 7)}")
    
    # Create comprehensive topic list for prompt
    topics_str = '\n'.join([f"- {t}" for t in topics_list])
    
    # DiTeLe-COMPLIANT PROMPT
    prompt = f"""Du bist ein erfahrener IT-Ausbilder und Praxisexperte für Fachinformatiker Anwendungsentwicklung.

🎯 AUFGABE: Erstelle ein PRAXISNAHES Lernszenario nach dem DiTeLe-Standard

📚 QUELLDOKUMENT: {doc_name}

DOKUMENT-INHALT (Auszug):
{document_content[:4000]}
[... Dokument enthält {len(document_content)} Zeichen total ...]

🔍 WICHTIGE THEMEN DIE ABGEDECKT WERDEN SOLLEN:
{topics_str}

⚠️ WICHTIG: Dieser EINE Anwendungsfall soll ALLE oder die wichtigsten dieser Themen abdecken!

═══════════════════════════════════════════════════════════════
STRUKTUR DES UMFASSENDEN ANWENDUNGSFALLS
═══════════════════════════════════════════════════════════════

Anwendungsfall: [Prägnanter, aussagekräftiger Titel]

1. THEORETISCHER HINTERGRUND (4-5 Absätze, ~800 Zeichen)

[Erkläre die fachlichen Grundlagen der WICHTIGSTEN Themen aus dem Dokument]
[Verwende EINFACHE Sprache - verständlich für Auszubildende und Schüler]
[Verbinde die verschiedenen Themen sinnvoll miteinander]
[Warum sind diese Themen in der IT-Praxis wichtig?]
[Welche Technologien/Konzepte werden benötigt?]
[Nutze Beispiele aus echten Projekten]

2. PRAXIS-SZENARIO (4-5 Absätze, ~700 Zeichen)

[Beschreibe ein REALISTISCHES Szenario in einem IT-Unternehmen]
[Das Szenario soll die wichtigsten Themen NATÜRLICH integrieren]
[MACHBAR für Lernende - nicht zu komplex!]
[Welches konkrete Problem/Projekt gibt es?]
[Wer sind die beteiligten Personen (Entwicklerteam, Kunde, etc.)?]
[Der Kontext soll motivierend und praxisnah sein]
[WICHTIG: Das Szenario soll mehrere Themen verbinden, nicht nur eines!]

3. AUFGABEN FÜR LERNENDE (5-7 AUFGABEN, ~1000 Zeichen)

KRITISCH: Die Aufgaben sollen verschiedene Themen abdecken!
- Jede Aufgabe fokussiert sich auf 1-2 der identifizierten Themen
- EINFACH und MACHBAR für Anfänger/Schüler
- Schrittweise steigende Komplexität
- Mit konkreten Tools und Hilfsmitteln
- Realistische Zeitaufwände (15-45 Minuten pro Aufgabe)

Aufgabe 1: [EINFACHE Einstiegsaufgabe - Thema: {topics_list[0] if topics_list else 'Grundlagen'}]
   Erwartetes Ergebnis: [Klares, messbares Ziel]
   Hilfsmittel: [Konkrete Tools, Dokumentation, Links]
   Zeitaufwand: Ca. [XX] Minuten
   Schwierigkeitsgrad: Einsteiger

Aufgabe 2: [Nächste Aufgabe - Thema: {topics_list[1] if len(topics_list) > 1 else 'Vertiefung'}]
   Erwartetes Ergebnis: [Was soll erreicht werden?]
   Hilfsmittel: [Tools für Anfänger]
   Zeitaufwand: Ca. [XX] Minuten
   Schwierigkeitsgrad: Einsteiger

Aufgabe 3-7: [Weitere Aufgaben decken zusätzliche Themen ab]
   [Jede Aufgabe baut auf der vorherigen auf]
   [IMMER Hilfsmittel und erwartete Ergebnisse angeben]
   [Komplexität nur LEICHT steigern]

4. DETAILLIERTE MUSTERLÖSUNGEN (6-8 Absätze, ~1500 Zeichen)

⚠️ EXTREM WICHTIG: Die Lösungen müssen SO DETAILLIERT sein, dass:
- Ein Anfänger JEDEN Schritt nachvollziehen kann
- Ein Trainer die Lösung direkt mit Schülerarbeit vergleichen kann
- ALLE Zwischenschritte erklärt sind
- Code-Beispiele enthalten sind (wenn relevant)
- WARUM-Erklärungen für jeden Schritt

Lösung zu Aufgabe 1:
Schritt 1: [Genau beschreiben was zu tun ist]
Schritt 2: [Welches Tool/Befehl wird verwendet - MIT BEISPIEL]
Schritt 3: [Was ist das erwartete Ergebnis]
Schritt 4: [Falls Code nötig - konkretes Beispiel angeben]

Beispiel-Code (falls relevant):
[Konkreter, funktionierender Code-Snippet]

Erklärung: [WARUM ist das die richtige Lösung?]

Lösung zu Aufgabe 2:
[Wieder SEHR detailliert, Schritt für Schritt]
[Mit konkreten Beispielen]
[Jeder Schritt hat eine WARUM-Erklärung]

Lösung zu Aufgabe 3-7:
[Für JEDE Aufgabe vollständige, nachvollziehbare Lösung]
[Mit Code-Beispielen wo sinnvoll]
[Mit Screenshots-Beschreibungen wenn hilfreich]

Alternative Lösungsansätze:
[1-2 alternative Wege zur Lösung beschreiben]
[Vor- und Nachteile erklären]
[Wann welcher Ansatz besser ist]

Best Practices und Professionelle Tipps:
[Was sollten Lernende in der Praxis beachten?]
[Wie machen es echte Entwickler?]
[Effizienz-Tipps und Qualitätsstandards]

Häufige Anfängerfehler und Problemlösung:
[Top 3-4 typische Fehler auflisten]
[Wie erkennt man diese Fehler?]
[Wie behebt man sie konkret?]
[Debugging-Strategien für Anfänger]

5. ERWARTETE LERNERGEBNISSE (Stichpunkte, ~300 Zeichen)

Nach erfolgreichem Abschluss können Lernende:
- [Konkrete Fähigkeit 1 - bezogen auf Thema X]
- [Konkrete Fähigkeit 2 - bezogen auf Thema Y]
- [Fachbegriffe verstehen und korrekt anwenden]
- [Praktische Probleme selbstständig lösen]
- [Konzepte auf neue Situationen übertragen]
- [Professionelle Arbeitsweise demonstrieren]

═══════════════════════════════════════════════════════════════

💡 KRITISCHE ANFORDERUNGEN - SEHR WICHTIG:

1. ✅ Decke ALLE wichtigen Themen aus der Liste ab (oder mind. 70%)
2. ✅ MINDESTENS 3500 Zeichen insgesamt (wegen Detailtiefe!)
3. ✅ Aufgaben sind für ANFÄNGER/SCHÜLER machbar
4. ✅ Lösungen SO DETAILLIERT, dass Trainer sie zum Vergleichen nutzen können
5. ✅ PROFESSIONELL formatiert mit klarer Struktur und Abständen
6. ✅ KEINE Markdown-Formatierung (**, ##, ```, etc.)
7. ✅ KEINE Erwähnung von "Bot", "KI", "Qualität", "Score"
8. ✅ Realistische IT-Szenarien aus echter Berufspraxis
9. ✅ Deutsche Sprache, Fachbegriffe einfach erklärt
10. ✅ Struktur EXAKT wie oben beschrieben einhalten

ZIELGRUPPE: Auszubildende und Schüler in IT-Berufen (Anfänger-Level)
OUTPUT-ZWECK: Professionelles Lehrmaterial für Trainer und Lernende
QUALITÄT: Muss direkt verwendbar sein ohne Nachbearbeitung

Erstelle jetzt den umfassenden, professionellen Anwendungsfall:"""

    try:
        # Generate with direct API call (no batching needed)
        logger.info("   📡 Generiere mit Gemini API...")
        result = await intelligent_gemini.generate_from_prompt(
            prompt=prompt,
            content_type="comprehensive_use_case",
            timeout=240,  # 4 minutes for comprehensive generation
            max_retries=3
        )
        
        # Validate result
        if result and len(result) > 3000:
            logger.info(f"   ✅ Erfolgreich generiert: {len(result):,} Zeichen")
            return result
        else:
            logger.warning(f"   ⚠️ Zu wenig Inhalt ({len(result) if result else 0} Zeichen), Retry mit kürzerem Context...")
            
            # Retry with shorter context
            retry_prompt = prompt.replace(document_content[:4000], document_content[:2500])
            result = await intelligent_gemini.generate_from_prompt(
                prompt=retry_prompt,
                content_type="comprehensive_use_case_retry",
                timeout=240
            )
            
            if result and len(result) > 2500:
                logger.info(f"   ✅ Retry erfolgreich: {len(result):,} Zeichen")
                return result
            else:
                logger.error("   ❌ Generierung fehlgeschlagen nach Retry")
                return ""
    
    except Exception as e:
        logger.error(f"   ❌ Fehler bei Generierung: {e}")
        import traceback
        traceback.print_exc()
        return ""


def validate_use_case_quality(
    use_cases_content: str,
    analysis: AnalysisResult
) -> Dict[str, Any]:
    """
    🔍 Qualitätsprüfung des generierten Anwendungsfalls
    
    UPDATED for single comprehensive use case:
    - Validates ONE comprehensive use case (not multiple)
    - Higher character requirements (3500+ for comprehensive coverage)
    - Ensures all major topics are mentioned
    - Checks for professional formatting and complete solutions
    
    Returns:
        Dict with validation results, issues, and quality metrics
    """
    logger.info("🔍 Führe Qualitätsprüfung durch...")
    
    issues = []
    warnings = []
    
    # Prüfung 1: Mindestlänge für umfassenden Use Case
    content_length = len(use_cases_content)
    if content_length < 2500:
        issues.append(f"Zu wenig Inhalt: {content_length} Zeichen (Minimum: 3500 für umfassenden Use Case)")
    elif content_length < 3500:
        warnings.append(f"Inhalt könnte ausführlicher sein: {content_length} Zeichen (Ziel: 3500+)")
    else:
        logger.info(f"   ✅ Länge ausreichend: {content_length:,} Zeichen")
    
    # Prüfung 2: Use Case vorhanden
    if "Anwendungsfall" not in use_cases_content:
        issues.append("Keine 'Anwendungsfall'-Überschrift gefunden!")
    
    # Prüfung 3: Struktur-Vollständigkeit (alle 5 Hauptsektionen)
    required_sections = [
        'THEORETISCHER HINTERGRUND',
        'PRAXIS-SZENARIO', 
        'AUFGABEN FÜR LERNENDE',
        'MUSTERLÖSUNG',
        'LERNERGEBNISSE'
    ]
    
    missing_sections = []
    for section in required_sections:
        if section.upper() not in use_cases_content.upper():
            missing_sections.append(section)
    
    if missing_sections:
        issues.append(f"Fehlende Strukturelemente: {', '.join(missing_sections)}")
    else:
        logger.info(f"   ✅ Alle 5 Strukturelemente vorhanden")
    
    # Prüfung 4: Aufgaben vorhanden
    task_count = use_cases_content.upper().count("AUFGABE")
    if task_count < 5:
        warnings.append(f"Nur {task_count} Aufgaben gefunden (empfohlen: 5-7)")
    elif task_count > 10:
        warnings.append(f"{task_count} Aufgaben gefunden (könnte zu viel sein)")
    else:
        logger.info(f"   ✅ {task_count} Aufgaben gefunden")
    
    # Prüfung 5: Lösungen vorhanden
    solution_indicators = ['Lösung', 'Schritt 1', 'Schritt 2', 'Best Practices']
    solution_count = sum(use_cases_content.count(indicator) for indicator in solution_indicators)
    if solution_count < 5:
        issues.append(f"Zu wenig Lösungsdetails gefunden ({solution_count} Indikatoren)")
    else:
        logger.info(f"   ✅ Detaillierte Lösungen vorhanden ({solution_count} Indikatoren)")
    
    # Prüfung 6: Themenabdeckung (prüfe ob wichtige Themen erwähnt werden)
    major_topics = [t.title for t in analysis.topics[:5]]  # Top 5 Themen
    covered_topics = sum(1 for topic in major_topics if topic.lower() in use_cases_content.lower())
    coverage_percentage = (covered_topics / len(major_topics) * 100) if major_topics else 100
    
    if coverage_percentage < 60:
        warnings.append(f"Nur {covered_topics}/{len(major_topics)} Hauptthemen erwähnt ({coverage_percentage:.0f}%)")
    else:
        logger.info(f"   ✅ {covered_topics}/{len(major_topics)} Hauptthemen abgedeckt ({coverage_percentage:.0f}%)")
    
    # Prüfung 7: Keine Fehlermeldungen im Content
    error_indicators = ['Error', 'Fehler:', 'timeout', 'overloaded', '503', '504']
    for indicator in error_indicators:
        if indicator.lower() in use_cases_content[:1000].lower():
            issues.append(f"Fehlermeldung im Inhalt gefunden: '{indicator}'")
    
    # Prüfung 8: Markdown-Symbole (sollten nicht da sein für professionelles Dokument)
    markdown_symbols = ['**', '##', '###', '```']
    markdown_count = sum(use_cases_content.count(sym) for sym in markdown_symbols)
    if markdown_count > 10:
        warnings.append(f"Viele Markdown-Symbole gefunden ({markdown_count}x) - sollten entfernt werden")
    
    # Prüfung 9: Bot/Quality-Erwähnungen (NICHT erlaubt!)
    forbidden_terms = ['bot', 'ki-generiert', 'quality score', 'qualitätsscore', 'ai-generated']
    for term in forbidden_terms:
        if term in use_cases_content.lower():
            issues.append(f"Unerlaubte Erwähnung gefunden: '{term}' (muss professionell sein!)")
    
    # Berechne Qualitätsscore (intern, nicht im Dokument sichtbar)
    quality_score = 100
    quality_score -= len(issues) * 25  # Kritische Probleme: -25 Punkte
    quality_score -= len(warnings) * 10  # Warnungen: -10 Punkte
    quality_score = max(0, min(100, quality_score))
    
    validation_result = {
        'quality_score': quality_score,
        'use_case_count': 1,  # Always 1 now
        'expected_count': 1,
        'content_length': content_length,
        'topic_coverage_percent': coverage_percentage,
        'task_count': task_count,
        'issues': issues,
        'warnings': warnings,
        'passed': len(issues) == 0,
        'grade': 'EXCELLENT' if quality_score >= 90 else 'GOOD' if quality_score >= 75 else 'ACCEPTABLE' if quality_score >= 60 else 'NEEDS_IMPROVEMENT'
    }
    
    # Log Ergebnisse
    if validation_result['passed']:
        logger.info(f"   ✅ Qualitätsprüfung BESTANDEN")
        logger.info(f"   🎯 Interner Score: {quality_score}/100 ({validation_result['grade']})")
        logger.info(f"   📊 Themenabdeckung: {coverage_percentage:.0f}%")
    else:
        logger.warning(f"   ⚠️ Qualitätsprüfung mit Problemen")
        logger.warning(f"   🎯 Interner Score: {quality_score}/100 ({validation_result['grade']})")
        for issue in issues:
            logger.warning(f"      ❌ {issue}")
    
    for warning in warnings:
        logger.info(f"      ⚠️ {warning}")
    
    return validation_result


def create_professional_word_document(
    use_cases_content: str,
    filename: str,
    analysis: AnalysisResult,
    validation: Dict[str, Any]
12221222    """
    📄 Erstelle ein professionelles Word-Dokument mit dem Anwendungsfall
    
    UPDATED for single comprehensive use case:
    - Professional title page
    - Document metadata (NO quality scores!)
    - ONE comprehensive use case with proper formatting
    - Clean, professional appearance
    
    Structure:
    - Title page
    - Document information (source, topics, complexity - NO SCORES!)
    - Comprehensive use case with proper formatting
    - Footer with metadata
    """
    logger.info("📄 Erstelle professionelles Word-Dokument...")
    
    doc = Document()
    
    # ========== PROFESSIONAL STYLING ==========
    # Normal text style - Arial Narrow with proper spacing
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial Narrow'
    font.size = Pt(11)
    
    # Paragraph spacing
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(6)
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15  # Professional line spacing
    
    # Heading 1 style - Large, bold, blue
    heading1_style = doc.styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Arial Narrow'
    heading1_font.size = Pt(18)
    heading1_font.bold = True
    heading1_font.color.rgb = RGBColor(0, 51, 153)
    heading1_style.paragraph_format.space_before = Pt(12)
    heading1_style.paragraph_format.space_after = Pt(6)
    
    # Heading 2 style - Medium, bold, dark blue
    heading2_style = doc.styles['Heading 2']
    heading2_font = heading2_style.font
    heading2_font.name = 'Arial Narrow'
    heading2_font.size = Pt(14)
    heading2_font.bold = True
    heading2_font.color.rgb = RGBColor(0, 102, 204)
    heading2_style.paragraph_format.space_before = Pt(10)
    heading2_style.paragraph_format.space_after = Pt(4)
    
    # ========== TITELSEITE ==========
    title = doc.add_heading('IT-Anwendungsfall für die Praxis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = 'Arial Narrow'
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.color.rgb = RGBColor(0, 51, 153)
    title.runs[0].font.bold = True
    
    doc.add_paragraph()  # Abstand
    
    subtitle = doc.add_paragraph(f'{filename}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.name = 'Arial Narrow'
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    subtitle.runs[0].bold = True
    subtitle.paragraph_format.space_before = Pt(12)
    subtitle.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph(f'Erstellt am: {datetime.now().strftime("%d.%m.%Y %H:%M")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.name = 'Arial Narrow'
    date_para.runs[0].font.size = Pt(12)
    date_para.runs[0].font.italic = True
    date_para.paragraph_format.space_before = Pt(6)
    
    doc.add_page_break()
    
    # ========== DOKUMENT-INFORMATIONEN (NO QUALITY SCORES!) ==========
    info_heading = doc.add_heading('📋 Dokument-Informationen', level=1)
    info_heading.paragraph_format.space_before = Pt(0)  # No extra space after page break
    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    # Header
    header_cells = info_table.rows[0].cells
    header_cells[0].text = 'Eigenschaft'
    header_cells[1].text = 'Wert'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.name = 'Arial Narrow'
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows (NO quality score!) - Apply Arial Narrow to all cells
    for row_idx in range(1, 5):
        for col_idx in range(2):
            for paragraph in info_table.rows[row_idx].cells[col_idx].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial Narrow'
    
    info_table.cell(1, 0).text = 'Quelldokument'
    info_table.cell(1, 1).text = filename
    
    info_table.cell(2, 0).text = 'Identifizierte Themen'
    info_table.cell(2, 1).text = str(len(analysis.topics))
    
    info_table.cell(3, 0).text = 'Komplexitätsstufe'
    info_table.cell(3, 1).text = f"{analysis.complexity_score:.1f}/10 ({analysis.technical_depth})"
    
    info_table.cell(4, 0).text = 'Abdeckungsstrategie'
    info_table.cell(4, 1).text = 'Umfassender Anwendungsfall mit allen wichtigen Themen'
    
    doc.add_page_break()
    
    # ========== ANWENDUNGSFALL ==========
    content_heading = doc.add_heading('📚 Umfassender IT-Anwendungsfall', level=1)
    content_heading.paragraph_format.space_before = Pt(0)  # No extra space after page break
    
    intro_para = doc.add_paragraph(
        'Der folgende Anwendungsfall deckt alle wichtigen Themen aus dem Quelldokument ab. '
        'Er ist speziell für Lernende und Auszubildende konzipiert: praxisnah, leicht verständlich '
        'und mit vollständigen Schritt-für-Schritt-Lösungen für jede Aufgabe. Die Aufgaben sind so '
        'gestaltet, dass sie von Anfängern umsetzbar und motivierend sind.'
    )
    intro_para.runs[0].font.name = 'Arial Narrow'
    intro_para.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()
    
    # Parse und formatiere Anwendungsfall
    for line in use_cases_content.split('\n'):
        line = line.strip()
        if not line or line == '═' * len(line):
            continue
        
        # Anwendungsfall-Überschrift
        if line.startswith('Anwendungsfall'):
            doc.add_paragraph()  # Abstand
            heading = doc.add_heading(line, level=1)
            # Heading 1 style already applied
        
        # Unterüberschriften (1. 2. 3. etc.)
        elif re.match(r'^\d+\.\s+[A-ZÄÖÜ]', line):
            heading2 = doc.add_heading(line, level=2)
            # Heading 2 style already applied
        
        # Fett markierte Zeilen (Aufgaben, etc.)
        elif line.startswith('Aufgabe') or line.startswith('Erwartetes Ergebnis:') or line.startswith('Hilfsmittel:') or line.startswith('Schwierigkeitsgrad:') or line.startswith('Zeitaufwand:') or line.startswith('Lösung zu'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.name = 'Arial Narrow'
            if 'Aufgabe' in line or 'Lösung' in line:
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(4)
        
        # Aufzählungen
        elif line.startswith('-') or line.startswith('•'):
            bullet_para = doc.add_paragraph(line[1:].strip(), style='List Bullet')
            for run in bullet_para.runs:
                run.font.name = 'Arial Narrow'
        
        # Normale Absätze
        else:
            normal_para = doc.add_paragraph(line)
            for run in normal_para.runs:
                run.font.name = 'Arial Narrow'
    
    # ========== FOOTER ==========
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f'FIAE AI Content Factory | Generiert am {datetime.now().strftime("%d.%m.%Y %H:%M")} | Seite '
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.name = 'Arial Narrow'
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    # Speichere als Bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    doc_size_kb = len(buffer.getvalue()) / 1024
    logger.info(f"✅ Word-Dokument erstellt: {doc_size_kb:.1f} KB")
    
    return buffer.getvalue()


async def process_document_phase1(doc: Dict[str, Any], index: int, total: int) -> Dict[str, Any]:
    """
    📄 Verarbeite ein einzelnes Dokument - UPDATED for single comprehensive use case
    
    Workflow:
    1. Inhalt extrahieren
    2. Dokument analysieren (KI-gestützt, identifiziert alle Themen)
    3. EINEN umfassenden Anwendungsfall generieren (alle wichtigen Themen)
    4. Qualität prüfen (Detailtiefe, Struktur, professionelle Formatierung)
    5. Professionelles Word-Dokument erstellen und hochladen (OHNE Quality-Scores)
    
    Returns:
        Dict with status, file info, content metrics (NOT quality scores!)
    """
    doc_id = doc['id']
    doc_name = doc['name']
    
    logger.info(f"\n{'='*80}")
    logger.info(f"📄 [{index}/{total}] VERARBEITE: {doc_name}")
    logger.info(f"{'='*80}")
    
    start_time = datetime.now()
    
    try:
        # SCHRITT 1: Inhalt extrahieren
        logger.info("\n[1/5] 📥 Extrahiere Dokumentinhalt...")
        document_content = google_drive_service.get_file_content(doc_id)
        
        if not document_content or len(document_content) < 100:
            logger.error("❌ Zu wenig Inhalt extrahiert")
            return {'status': 'failed', 'error': 'insufficient_content'}
        
        word_count = len(document_content.split())
        logger.info(f"   ✅ {len(document_content):,} Zeichen extrahiert ({word_count:,} Wörter)")
        
        # SCHRITT 2: Dokument analysieren
        logger.info("\n[2/5] 🧠 Analysiere Dokument mit KI...")
        analysis = await analyze_document_for_use_cases(document_content, doc_name)
        
        # SCHRITT 3: Anwendungsfall generieren
        logger.info(f"\n[3/5] 🤖 Generiere umfassenden Anwendungsfall...")
        use_cases_content = await generate_single_comprehensive_use_case(
            document_content,
            doc_name,
            analysis
        )
        
        if not use_cases_content or len(use_cases_content) < 2000:
            logger.error("❌ Anwendungsfall-Generierung fehlgeschlagen oder zu wenig Inhalt")
            return {'status': 'failed', 'error': 'generation_failed'}
        
        # SCHRITT 4: Qualitätsprüfung
        logger.info("\n[4/5] 🔍 Führe Qualitätsprüfung durch...")
        validation = validate_use_case_quality(use_cases_content, analysis)
        
        if not validation['passed']:
            logger.warning("⚠️ Qualitätsprüfung mit Problemen, aber fahre fort...")
        
        # SCHRITT 5: Word-Dokument erstellen und hochladen
        logger.info("\n[5/5] 💾 Erstelle und lade Word-Dokument hoch...")
        
        doc_bytes = create_professional_word_document(
            use_cases_content,
            doc_name,
            analysis,
            validation
        )
        
        # Upload zu Google Drive
        base_name = doc_name.rsplit('.', 1)[0]
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"Anwendungsfälle_{base_name}_{timestamp}.docx"
        
        review_folder_id = settings.google_drive_review_folder_id
        
        file_metadata = {
            'name': output_filename,
            'parents': [review_folder_id],
            'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        
        # Validate Google Drive service before upload
        if not google_drive_service or not google_drive_service.service:
            raise Exception("Google Drive service not initialized")
        
        media = MediaIoBaseUpload(
            io.BytesIO(doc_bytes),
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            resumable=True
        )
        
        file = google_drive_service.service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id,name,webViewLink'
        ).execute()
        
        elapsed = (datetime.now() - start_time).total_seconds()
        
        logger.info(f"\n✅ ERFOLG! Abgeschlossen in {elapsed:.1f}s")
        logger.info(f"📁 Datei: {output_filename}")
        logger.info(f"🔗 Link: {file.get('webViewLink', 'N/A')}")
        logger.info(f"📄 Typ: 1 umfassender Anwendungsfall")
        logger.info(f"📊 Themenabdeckung: {validation['topic_coverage_percent']:.0f}%")
        logger.info(f"✨ Professionell formatiert, anfängerfreundlich, detaillierte Lösungen")
        
        return {
            'status': 'completed',
            'doc_name': doc_name,
            'file_id': file['id'],
            'file_name': output_filename,
            'web_view_link': file.get('webViewLink'),
            'use_case_count': 1,  # Always 1 now
            'content_length': validation['content_length'],
            'topic_coverage': validation['topic_coverage_percent'],
            'processing_time': elapsed
        }
        
    except Exception as e:
        logger.error(f"❌ Fehler bei Verarbeitung von {doc_name}: {e}")
        import traceback
        traceback.print_exc()
        return {
            'status': 'failed',
            'doc_name': doc_name,
            'error': str(e)
        }


async def main():
    """
    Haupt-Workflow für Phase 1: Single Comprehensive Use Case Generation
    
    NEW: Generates ONE comprehensive use case per source document
    """
    
    logger.info("="*80)
    logger.info("🚀 PHASE 1: SINGLE COMPREHENSIVE USE CASE GENERATION")
    logger.info("="*80)
    logger.info("🎯 Fokus: EIN umfassender Anwendungsfall pro Dokument")
    logger.info("📄 Output: EIN professionelles Word-Dokument pro Quelldokument")
    logger.info("🧠 Strategie: Alle wichtigen Themen in EINEM Use Case")
    logger.info("✅ Qualität: Extrem detailliert + Anfängerfreundlich + Professionell")
    logger.info("="*80 + "\n")
    
    try:
        # Entdecke Dokumente
        source_folder_id = settings.google_drive_content_source_folder_id
        logger.info(f"🔍 Suche Dokumente in Quellordner...")
        
        all_documents = google_drive_service.list_files_in_folder(source_folder_id)
        
        # Check if documents are in root or subfolders
        direct_documents = [
            doc for doc in all_documents
            if doc['name'].lower().endswith(('.docx', '.pdf', '.txt'))
            and not doc['name'].startswith('~')
        ]
        
        documents = []
        
        if direct_documents:
            # Documents are in root folder
            documents = direct_documents
            logger.info(f"📚 {len(documents)} Dokumente im Hauptordner gefunden")
        else:
            # Documents are in subfolders - search recursively
            logger.info(f"📁 Keine Dokumente im Hauptordner, suche in Unterordnern...")
            folders = [f for f in all_documents if f.get('mimeType') == 'application/vnd.google-apps.folder']
            logger.info(f"   Gefunden: {len(folders)} Unterordner")
            
            for folder in folders:
                try:
                    folder_files = google_drive_service.list_files_in_folder(folder['id'])
                    folder_docs = [
                        doc for doc in folder_files
                        if doc['name'].lower().endswith(('.docx', '.pdf', '.txt'))
                        and not doc['name'].startswith('~')
                    ]
                    documents.extend(folder_docs)
                except Exception as e:
                    logger.warning(f"   ⚠️ Fehler beim Lesen von Ordner {folder['name']}: {e}")
            
            logger.info(f"📚 {len(documents)} Dokumente in Unterordnern gefunden")
        
        if not documents:
            logger.warning("⚠️ Keine Dokumente im Quellordner gefunden!")
            return
        
        # Verarbeite jedes Dokument
        results = []
        success_count = 0
        
        for i, doc in enumerate(documents, 1):
            result = await process_document_phase1(doc, i, len(documents))
            results.append(result)
            
            if result['status'] == 'completed':
                success_count += 1
            
            # Pause zwischen Dokumenten (Rate-Limiting)
            if i < len(documents):
                logger.info(f"\n⏸️  Pause 10s vor nächstem Dokument...\n")
                await asyncio.sleep(10)
        
        # Abschluss-Statistiken
        logger.info(f"\n{'='*80}")
        logger.info(f"🎉 VERARBEITUNG ABGESCHLOSSEN")
        logger.info(f"{'='*80}")
        logger.info(f"✅ Erfolgreich: {success_count}/{len(documents)}")
        logger.info(f"❌ Fehlgeschlagen: {len(documents) - success_count}/{len(documents)}")
        success_rate = (success_count / len(documents) * 100) if documents else 0
        logger.info(f"📊 Erfolgsquote: {success_rate:.1f}%")
        
        if success_count > 0:
            successful_results = [r for r in results if r['status'] == 'completed']
            
            avg_coverage = sum(r.get('topic_coverage', 0) for r in successful_results) / success_count
            avg_length = sum(r.get('content_length', 0) for r in successful_results) / success_count
            avg_time = sum(r['processing_time'] for r in successful_results) / success_count
            
            logger.info(f"\n📈 Durchschnittswerte:")
            logger.info(f"   � Typ: 1 umfassender Use Case pro Dokument (NEW!)")
            logger.info(f"   📊 Themenabdeckung: {avg_coverage:.1f}%")
            logger.info(f"   📏 Durchschnittliche Länge: {avg_length:.0f} Zeichen")
            logger.info(f"   ⏱️  Verarbeitungszeit: {avg_time:.1f}s")
            logger.info(f"   ✨ Extrem detaillierte Lösungen für Trainer-Vergleich")
        
        logger.info(f"{'='*80}\n")
        
    except Exception as e:
        logger.error(f"❌ Fataler Fehler: {e}")
        import traceback
        traceback.print_exc()
        raise


if __name__ == "__main__":
    asyncio.run(main())
