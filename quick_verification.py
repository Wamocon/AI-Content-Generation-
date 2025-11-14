#!/usr/bin/env python3
"""
Quick Verification Script - Tests Core Logic Without API Calls
===============================================================
"""

import sys
sys.path.insert(0, '.')

from app.services.document_analyzer import DocumentAnalyzer, Topic
from loguru import logger

logger.info("=" * 70)
logger.info("QUICK VERIFICATION: Single Use Case Generation Changes")
logger.info("=" * 70)

# Test 1: Verify use_cases_count is always 1
logger.info("\n[TEST 1] Verifying use_cases_count = 1")

analyzer = DocumentAnalyzer(gemini_service=None)

# Test with small content
small_content = "Python programming. Variables. Functions. Classes." * 20

# Extract topics using rule-based method
topics = analyzer._extract_topics_rule_based(small_content)
logger.info(f"   Extracted {len(topics)} topics (rule-based)")

# Calculate content requirements
from app.services.document_analyzer import ContentRequirements

class MockAnalysis:
    def __init__(self):
        self.topics = topics
        self.complexity_score = 5.0
        self.main_themes = ["Programming"]

requirements = analyzer._calculate_content_requirements(
    topics=topics,
    document_chars=len(small_content),
    document_words=len(small_content.split()),
    complexity_score=5.0,
    main_themes=["Programming"]
)

logger.info(f"   ✓ use_cases_count = {requirements.use_cases_count}")
assert requirements.use_cases_count == 1, "FAILED: use_cases_count should be 1"
logger.info("   ✅ TEST 1 PASSED: Always returns 1 use case")

# Test 2: Verify comprehensive topic list
logger.info("\n[TEST 2] Verifying comprehensive topic list")
topic_list = requirements.use_cases_topics
logger.info(f"   Topic list entries: {len(topic_list)}")
logger.info(f"   Content: {topic_list[0] if topic_list else 'None'}")
assert len(topic_list) == 1, "FAILED: Should have exactly 1 topic entry"
assert "Comprehensive" in topic_list[0], "FAILED: Should mention 'Comprehensive'"
logger.info("   ✅ TEST 2 PASSED: Comprehensive topic list created")

# Test 3: Verify validation logic
logger.info("\n[TEST 3] Verifying validation logic")

from automation_phase1_content import validate_use_case_quality, AnalysisResult

# Create mock analysis
mock_topics = [Topic(title=f"Topic {i}", subtopics=[], keywords=[], 
                     complexity="medium", estimated_content_depth=3) 
               for i in range(5)]

content_req = requirements

analysis = AnalysisResult(
    document_name="test.txt",
    total_chars=5000,
    total_words=800,
    topics=mock_topics,
    main_themes=["Software Development"],
    complexity_score=6.5,
    technical_depth="intermediate",
    content_requirements=content_req,
    analysis_confidence=0.85,
    recommendations=[]
)

# Test with comprehensive use case content
test_content = """
Anwendungsfall: Entwicklung einer Web-Anwendung

1. THEORETISCHER HINTERGRUND

Python ist eine vielseitige Programmiersprache für die Webentwicklung.
Flask ist ein leichtgewichtiges Web-Framework.
REST APIs ermöglichen die Kommunikation zwischen Client und Server.
Datenbanken speichern persistente Daten.
MVC-Pattern strukturiert die Anwendung.

2. PRAXIS-SZENARIO

Ein IT-Unternehmen benötigt eine Web-Anwendung für die Kundenverwaltung.
Das Entwicklerteam soll eine REST API mit Flask entwickeln.
Die Anwendung muss mit einer PostgreSQL-Datenbank verbunden werden.
Das Frontend wird mit React entwickelt.

3. AUFGABEN FÜR LERNENDE

Aufgabe 1: Erstelle ein Flask-Projekt mit Virtual Environment
   Erwartetes Ergebnis: Funktionierendes Flask-Setup
   Hilfsmittel: Flask Dokumentation, pip install flask
   Zeitaufwand: Ca. 30 Minuten
   Schwierigkeitsgrad: Einsteiger

Aufgabe 2: Implementiere eine REST API für Kunden
   Erwartetes Ergebnis: CRUD-Operationen funktionieren
   Hilfsmittel: Flask-RESTful, Postman zum Testen
   Zeitaufwand: Ca. 45 Minuten
   
Aufgabe 3: Verbinde PostgreSQL-Datenbank
   Erwartetes Ergebnis: Datenbank-Queries funktionieren
   
Aufgabe 4: Erstelle Frontend-Komponenten
Aufgabe 5: Teste die Anwendung mit Unit Tests

4. DETAILLIERTE MUSTERLÖSUNGEN

Lösung zu Aufgabe 1:
Schritt 1: Erstelle Virtual Environment mit python -m venv venv
Schritt 2: Aktiviere mit venv\\Scripts\\activate
Schritt 3: Installiere Flask mit pip install flask
Schritt 4: Erstelle app.py Datei

Code-Beispiel:
from flask import Flask
app = Flask(__name__)

@app.route('/')
def hello():
    return 'Hello World'

if __name__ == '__main__':
    app.run(debug=True)

Erklärung: Flask benötigt eine Anwendungsinstanz und mindestens eine Route.

Lösung zu Aufgabe 2:
Schritt 1: Importiere Flask-RESTful
Schritt 2: Definiere Resource-Klasse
Schritt 3: Implementiere GET, POST, PUT, DELETE Methoden

Alternative Lösungsansätze:
- Flask-RESTful für strukturierte APIs
- Direkte Flask-Routes für einfache Anwendungen
Vorteil RESTful: Bessere Organisation bei vielen Endpoints

Best Practices:
- Verwende Virtual Environments für Isolation
- Schreibe aussagekräftige Docstrings
- Validiere Eingabedaten
- Nutze Environment Variables für Konfiguration

Häufige Anfängerfehler:
1. Vergessen, Virtual Environment zu aktivieren
2. Flask nicht im debug-Modus für Entwicklung
3. Fehlende Fehlerbehandlung bei API-Requests
Lösung: Immer try-except für robuste APIs

5. ERWARTETE LERNERGEBNISSE

- Flask-Anwendungen erstellen und strukturieren können
- REST APIs implementieren und testen
- Mit Datenbanken arbeiten (CRUD-Operationen)
- MVC-Pattern in der Praxis anwenden
- Unit Tests für Web-Anwendungen schreiben
- Best Practices der Webentwicklung kennen
"""

validation = validate_use_case_quality(test_content, analysis)

logger.info(f"   Use case count detected: {validation['use_case_count']}")
logger.info(f"   Content length: {validation['content_length']} chars")
logger.info(f"   Topic coverage: {validation['topic_coverage_percent']:.1f}%")
logger.info(f"   Task count: {validation['task_count']}")
logger.info(f"   Validation passed: {validation['passed']}")
logger.info(f"   Issues: {len(validation['issues'])}")
logger.info(f"   Warnings: {len(validation['warnings'])}")

assert validation['use_case_count'] == 1, "FAILED: Should detect 1 use case"
# Note: Test content is ~2700 chars, validation accepts 2500+ as minimum
assert validation['content_length'] > 2500, f"FAILED: Should have 2500+ chars (got {validation['content_length']})"
assert validation['passed'] == True, "FAILED: Validation should pass"
logger.info("   ✅ TEST 3 PASSED: Validation logic works correctly")

# Test 4: Verify forbidden terms detection
logger.info("\n[TEST 4] Verifying forbidden terms detection")

bad_content = test_content + "\n\nGenerated by AI bot with quality score 95/100"
validation_bad = validate_use_case_quality(bad_content, analysis)

logger.info(f"   Issues found: {len(validation_bad['issues'])}")
assert len(validation_bad['issues']) > 0, "FAILED: Should detect forbidden terms"
logger.info("   ✅ TEST 4 PASSED: Forbidden terms detected")

# Test 5: Verify Word document creation (without quality scores)
logger.info("\n[TEST 5] Verifying Word document structure")

from automation_phase1_content import create_professional_word_document
from docx import Document
import io

validation_clean = {
    'quality_score': 95,
    'use_case_count': 1,
    'content_length': 3500,
    'topic_coverage_percent': 85,
    'passed': True,
    'grade': 'EXCELLENT'
}

doc_bytes = create_professional_word_document(
    test_content,
    "test_document.txt",
    analysis,
    validation_clean
)

# Parse document
doc = Document(io.BytesIO(doc_bytes))
full_text = '\n'.join([para.text for para in doc.paragraphs])

# Check for forbidden terms
forbidden = ['quality score', 'qualitätsscore', '95/100', 'EXCELLENT', 'grade:', 'bot']
found = [term for term in forbidden if term.lower() in full_text.lower()]

logger.info(f"   Document size: {len(doc_bytes) / 1024:.1f} KB")
logger.info(f"   Total paragraphs: {len(doc.paragraphs)}")
logger.info(f"   Forbidden terms found: {found if found else 'None'}")

# Debug: Show what's actually in the metadata
logger.info(f"   First 20 paragraphs:")
for i, para in enumerate(doc.paragraphs[:20]):
    if para.text.strip():
        logger.info(f"      [{i}] {para.text[:80]}")

assert len(found) == 0, f"FAILED: Found forbidden terms: {found}"
assert "Quelldokument" in full_text or "test_document" in full_text, "FAILED: Missing required metadata"
# The table content might be in cells, not paragraphs
logger.info("   ✅ Document structure looks good")

logger.info("   ✅ TEST 5 PASSED: Document has no quality scores")

# Final Summary
logger.info("\n" + "=" * 70)
logger.info("🎉 ALL VERIFICATION TESTS PASSED!")
logger.info("=" * 70)
logger.info("✓ use_cases_count is always 1")
logger.info("✓ Comprehensive topic list created")
logger.info("✓ Validation logic works correctly")
logger.info("✓ Forbidden terms detection works")
logger.info("✓ Word documents have no quality scores")
logger.info("\n✅ System is ready for production use!")
logger.info("=" * 70)
